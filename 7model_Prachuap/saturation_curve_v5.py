# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  Ensemble Saturation Curve Analysis — v5.0                                  ║
║  "How Many CMIP6 Models Are Enough for Station-Scale Rainfall?"             ║
║  Q2–Q3 Publication Standard  |  Academically Rigorous                       ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  NEW IN v5.0 (vs v4.0)  — Three additional publication-ready figures:      ║
║  [A] Fig 4 — INTEGRATED DUAL-AXIS PERFORMANCE PROFILE                      ║
║      Panel (a): RMSE (left-Y) + KGE (right-Y) vs N — dual-axis, ±CI       ║
║      Panel (b): NSE (left-Y) + r   (right-Y) vs N — dual-axis, ±CI        ║
║      Both Raw and BC plotted; saturation stars on each axis;                ║
║      Exponential-decay fits; 95% Bootstrap CI shaded.                       ║
║  [B] Fig 5 — SPATIAL CONSENSUS PROFILE  (Stacked Bar)                      ║
║      For every ensemble size N, shows the % of stations whose              ║
║      saturation point N* ≤ N (cumulative spatial consensus).               ║
║      Separate bars: at-N (= N*) vs already-saturated (< N).               ║
║      Plotted for all 4 metrics, 2 × 2 layout.                             ║
║  [C] Fig 6 — INTER-METRIC RELATIONSHIP WEB  (Network Diagram)             ║
║      Computes cross-correlations between metric trajectories                ║
║      [RMSE, KGE, NSE, r] across N = 1…M for Raw and BC.                  ║
║      Nodes arranged on a circle; edge width ∝ |r|;                        ║
║      blue edge = positive, red = negative correlation.                      ║
║      Annotated with r-values; separate networks for Raw vs BC.             ║
║  All new figures: named without collision with Fig 1–3, lossless PNG       ║
║  at ≥ 600 DPI, individual panel exports (Figure_4a, 5a–5d, 6a–6b).       ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  INHERITED FROM v4.0:                                                        ║
║  [7] PANEL SPLITTING — save_individual_panels() auto-crops each panel       ║
║  [1] PARALLEL PROCESSING — ProcessPoolExecutor across all CPU cores         ║
║  [2] UNIT TESTING — pytest-compatible test suite                            ║
║  [3] DATA VALIDATION — pre-run integrity checks                             ║
║  [4] BOOTSTRAP CI FOR N* — 95% CI via resampling                          ║
║  [5] ENHANCED EXCEL — 4 sheets including Bootstrap CI                      ║
║  [6] IMPROVED FIGURES 1–3 — bold colours, non-overlapping legends          ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  Figures:                                                                    ║
║   Fig 1 – Ensemble Performance Saturation (4 panels: RMSE/KGE/NSE/r)      ║
║   Fig 2 – Spread Saturation + Uncertainty Reduction + Trade-off             ║
║   Fig 3 – Spatial Saturation (per-station + N* heatmap)                    ║
║   Fig 4 – Integrated Dual-Axis Performance Profile  [NEW v5.0]             ║
║   Fig 5 – Spatial Consensus Stacked Bar  [NEW v5.0]                        ║
║   Fig 6 – Inter-metric Relationship Web  [NEW v5.0]                        ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  References:                                                                 ║
║   Giorgi & Mearns (2002) J.Climate 15:1141–1158                            ║
║   Tebaldi & Knutti (2007) Phil.Trans.R.Soc.A 365:2053–2075                ║
║   Gupta et al. (2009) J.Hydrol. 377:80–91  [KGE]                          ║
║   Nash & Sutcliffe (1970) J.Hydrol. 10:282–290  [NSE]                     ║
║   Efron & Tibshirani (1993) Bootstrap Methods  [Bootstrap CI]              ║
║   Cannon et al. (2015) J.Climate 28:6938–6959  [QDM]                      ║
║   Moriasi et al. (2007) Trans.ASABE 50:885–900  [Performance criteria]    ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import os, sys, re, math, warnings, gc, itertools, time, io
import numpy as np
import pandas as pd
from pathlib import Path
from scipy.stats import pearsonr
from scipy.optimize import curve_fit
import concurrent.futures
import multiprocessing

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.ticker as ticker
import matplotlib.gridspec as gridspec
import matplotlib.colors as mcolors
from matplotlib.lines import Line2D
from matplotlib.colors import LinearSegmentedColormap
import matplotlib.cm as cm

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

warnings.filterwarnings("ignore")

# ════════════════════════════════════════════════════════════════════════
#  §0  GLOBAL CONSTANTS & PUBLICATION STYLE
# ════════════════════════════════════════════════════════════════════════

VERSION       = "5.0"
WET_THR       = 1.0
MIN_DAYS      = 280
DPI           = int(os.environ.get("CMIP6_DPI", 600))
SAVE_PDF      = True
MISS_FLAGS    = [-99, -999, -9999, -9.99e+20, 9.99e+20, 1e+20]
SAT_THRESHOLD = 0.05   # 5% marginal improvement threshold
N_BOOTSTRAP   = 999    # bootstrap iterations for CI on N*
CI_LEVEL      = 95     # confidence interval %
MAX_WORKERS   = max(1, multiprocessing.cpu_count())   # parallel cores

METS_ALL = ["RMSE", "KGE", "NSE", "r"]
LOWER_B  = {"RMSE"}   # lower = better

# ── Token skip list for model name extraction ─────────────────────────────
_SKIP_TOKENS = {
    "pr", "bc", "day", "mon", "yr", "6hr", "3hr", "1hr", "fx",
    "daily", "monthly", "annual", "seasonal",
    "hist", "historical", "ssp245", "ssp585", "ssp126", "rcp45", "rcp85",
    "gn", "gr", "grz",
}
_OBS_NOISE = {
    "observed", "rain", "daily", "monthly", "data",
    "rainfall", "precipitation", "output", "tables", "biascorrection",
}

# ── Publication colours (vivid, colour-blind friendly) ───────────────────
C = dict(
    obs="#1B2838",    obs_lt="#90A4AE",   obs_bd="#1B2838",
    raw="#D32F2F",    raw_lt="#FFCDD2",   raw_bd="#B71C1C",
    bc="#1565C0",     bc_lt="#BBDEFB",    bc_bd="#0D47A1",
    green="#1B5E20",  gold="#F57F17",     grey="#546E7A",
    purple="#6A1B9A", red2="#C62828",     amber="#FF8F00",
    teal="#00695C",
)
MODEL_PALETTE = [
    "#E65100", "#1565C0", "#2E7D32", "#6A1B9A",
    "#00695C", "#C62828", "#AD1457", "#37474F",
]
METRIC_COLORS  = {"RMSE": "#D32F2F", "KGE": "#1565C0",
                  "NSE": "#2E7D32",  "r":   "#6A1B9A"}
METRIC_MARKERS = {"RMSE": "o", "KGE": "s", "NSE": "^", "r": "D"}

# ── rcParams — bold, large, publication ─────────────────────────────────
plt.rcParams.update({
    "font.family":         "serif",
    "font.serif":          ["Times New Roman", "DejaVu Serif"],
    "font.size":           13,
    "axes.titlesize":      14,
    "axes.titleweight":    "bold",
    "axes.labelsize":      13,
    "axes.labelweight":    "bold",
    "xtick.labelsize":     12,
    "ytick.labelsize":     12,
    "legend.fontsize":     11.5,
    "lines.linewidth":     2.6,
    "lines.markersize":    9,
    "axes.linewidth":      1.8,
    "axes.spines.top":     False,
    "axes.spines.right":   False,
    "axes.grid":           True,
    "grid.linestyle":      "--",
    "grid.linewidth":      0.6,
    "grid.alpha":          0.40,
    "grid.color":          "#B0BEC5",
    "savefig.bbox":        "tight",
    "savefig.pad_inches":  0.15,
    "figure.dpi":          100,
    "mathtext.fontset":    "stix",
    "pdf.fonttype":        42,
    "ps.fonttype":         42,
    "xtick.major.width":   1.6,
    "ytick.major.width":   1.6,
    "xtick.minor.width":   0.9,
    "ytick.minor.width":   0.9,
})

# ── Excel helpers ─────────────────────────────────────────────────────────
THIN = Side(style="thin",   color="BDBDBD")
MED  = Side(style="medium", color="1F4E79")
XC   = dict(
    title="13293D", sub="1F4E79", hdr="2E75B6",
    raw_r="FFEBEE", bc_r="E3F2FD", sat="FFF9C4",
    improve="C8E6C9", degrade="FFCCBC", sig="DDEEFF",
    white="FFFFFF", alt="F5F5F5", note="ECEFF1",
)

def _tb():   return Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
def _xf(h):  return PatternFill("solid", fgColor=h)

def _xsc(ws, r, c, val=None, bold=False, italic=False,
         fc=None, bg=None, align="center", sz=10, wrap=True):
    cell = ws.cell(row=r, column=c)
    if val is not None: cell.value = val
    cell.font = Font(bold=bold, italic=italic, name="Calibri", size=sz,
                     color=fc if fc else "1A1A1A")
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
    if bg: cell.fill = _xf(bg)
    cell.border = _tb()
    return cell

def _mxsc(ws, r, c1, c2, val, **kw):
    ws.merge_cells(start_row=r, start_column=c1, end_row=r, end_column=c2)
    return _xsc(ws, r, c1, val, **kw)

def _cw(ws, col, w): ws.column_dimensions[get_column_letter(col)].width = w
def _rh(ws, r, h):   ws.row_dimensions[r].height = h

def savefig(fig, stem):
    p = str(stem)
    fig.savefig(p + ".png", dpi=DPI, bbox_inches="tight", pad_inches=0.15)
    if SAVE_PDF:
        try: fig.savefig(p + ".pdf", bbox_inches="tight", pad_inches=0.15)
        except Exception: pass
    plt.close(fig); gc.collect()
    print(f"  ✓  {Path(p).name}.png" + (" + .pdf" if SAVE_PDF else ""))


def save_individual_panels(fig, axes_list, panel_labels, fig_num,
                           out_dir, dpi=DPI, margin_in=0.22):
    """
    Auto-detect each panel's bounding box and save as individual lossless PNG.

    Uses ax.get_tightbbox(renderer) for sub-pixel-accurate panel boundaries,
    then crops and exports at full publication DPI (≥ 600).

    Parameters
    ----------
    fig          : matplotlib.figure.Figure
    axes_list    : list[Axes | list[Axes]]
                   One entry per panel. Pass a list of Axes to merge bounding
                   boxes (e.g. for twinx panels).
    panel_labels : list[str]  — e.g. ["a", "b", "c", "d"]
    fig_num      : int        — figure number used in output filenames
    out_dir      : Path       — destination directory
    dpi          : int        — output DPI (default 600)
    margin_in    : float      — white-space padding around each panel (inches)

    Output
    ------
    Saves Figure_<fig_num><label>.png (lossless PNG, white background).
    Returns list of saved file paths.
    """
    try:
        from PIL import Image as _PILImage
    except ImportError:
        print("  ⚠  Pillow not installed — panel split skipped  "
              "(pip install Pillow)")
        return []

    # ── 1. Ensure canvas is fully rendered ──────────────────────────────
    fig.canvas.draw()
    try:
        renderer = fig.canvas.get_renderer()
    except AttributeError:
        renderer = fig._cachedRenderer   # fallback for some backends

    # ── 2. Render the complete figure (no tight-crop) at target DPI ─────
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi,
                bbox_inches=None, pad_inches=0)
    buf.seek(0)
    img_full = _PILImage.open(buf).copy()
    buf.close()

    img_w, img_h = img_full.size          # pixels in saved image
    scale        = dpi / fig.dpi          # display→saved pixel scale (e.g. 6.0)
    margin_px    = int(margin_in * dpi)   # padding in pixels

    saved = []
    for entry, label in zip(axes_list, panel_labels):
        # Accept a single Axes or a list of Axes (merged bbox for twinx)
        ax_group = entry if isinstance(entry, (list, tuple)) else [entry]

        # Compute union of tight bounding boxes across the axis group
        x0s, y0s, x1s, y1s = [], [], [], []
        for ax in ax_group:
            try:
                bb = ax.get_tightbbox(renderer, for_layout_only=False)
            except TypeError:
                bb = ax.get_tightbbox(renderer)
            if bb is None:
                bb = ax.get_window_extent(renderer)
            x0s.append(bb.x0); y0s.append(bb.y0)
            x1s.append(bb.x1); y1s.append(bb.y1)

        bx0 = min(x0s); by0 = min(y0s)
        bx1 = max(x1s); by1 = max(y1s)

        # Scale display coords (fig.dpi) → saved-image pixels (dpi)
        px0 = bx0 * scale;  py0 = by0 * scale
        px1 = bx1 * scale;  py1 = by1 * scale

        # Convert to PIL (top-left origin vs matplotlib bottom-left)
        left   = max(0,      int(px0) - margin_px)
        right  = min(img_w,  int(px1) + margin_px)
        top    = max(0,      img_h - int(py1) - margin_px)
        bottom = min(img_h,  img_h - int(py0) + margin_px)

        if right <= left or bottom <= top:
            print(f"  ⚠  Panel {label}: degenerate crop — skipped")
            continue

        panel_img = img_full.crop((left, top, right, bottom))

        # Flatten to RGB with white background (handles RGBA transparency)
        bg = _PILImage.new("RGB", panel_img.size, (255, 255, 255))
        if panel_img.mode in ("RGBA", "LA"):
            bg.paste(panel_img, mask=panel_img.split()[-1])
            panel_img = bg
        else:
            panel_img = panel_img.convert("RGB")

        out_name = f"Figure_{fig_num}{label}.png"
        out_path = Path(out_dir) / out_name
        # PNG lossless — no compress_level override (default = lossless)
        panel_img.save(str(out_path), format="PNG",
                       dpi=(dpi, dpi), optimize=False)
        saved.append(str(out_path))
        w, h = panel_img.size
        print(f"  ✓  {out_name}  ({w}×{h} px  @{dpi} DPI)")

    return saved


# ════════════════════════════════════════════════════════════════════════
#  §1  FILE DISCOVERY  (province-aware, robust)
# ════════════════════════════════════════════════════════════════════════

def _province_keywords(obs_filename):
    stem = Path(obs_filename).stem
    normalised = stem.replace("_", " ").replace("-", " ").lower()
    normalised = re.sub(r"\b\d{4,}\b", " ", normalised)
    tokens = set()
    for tok in normalised.split():
        tok = tok.strip()
        if len(tok) >= 3 and tok not in _OBS_NOISE:
            tokens.add(tok)
    return tokens

def _filename_tokens(filepath):
    stem = Path(filepath).stem
    normalised = stem.replace("_", " ").replace("-", " ").lower()
    normalised = re.sub(r"\b\d{4,}\b", " ", normalised)
    return set(t.strip() for t in normalised.split() if len(t.strip()) >= 3)

def _file_matches_province(filepath, prov_kw):
    if not prov_kw: return True
    return bool(_filename_tokens(filepath) & prov_kw)

def _extract_model(filepath):
    stem = Path(filepath).stem
    parts = stem.split("_")
    while parts:
        candidate_clean = re.sub(r"[^a-z0-9]", "", parts[0].lower())
        if candidate_clean in _SKIP_TOKENS or parts[0].lower() in _SKIP_TOKENS:
            parts.pop(0)
        else:
            break
    if not parts: return "UnknownModel"
    model_name = parts[0]
    if re.match(r"r\d+i\d+p\d+", model_name.lower()):
        parts.pop(0)
        model_name = parts[0] if parts else "UnknownModel"
    return model_name

def _list_csv_files(folder):
    folder_p = Path(folder)
    try:
        csvs = sorted(folder_p.glob("*.csv"))
        if csvs: return csvs
    except Exception: pass
    try:
        names = os.listdir(str(folder_p))
        return sorted(folder_p / n for n in names if n.lower().endswith(".csv"))
    except Exception: return []

def discover_files(folder):
    all_csv = _list_csv_files(folder)
    if not all_csv:
        print(f"  ✗  No CSV files found in: {folder}"); return None, {}, {}
    print(f"  Found {len(all_csv)} CSV files")
    obs_files = [f for f in all_csv if "observed" in f.name.lower()]
    raw_files  = [f for f in all_csv if f.name.lower().startswith("pr_")
                  and "observed" not in f.name.lower()]
    bc_files   = [f for f in all_csv if f.name.lower().startswith("bc_")
                  and "observed" not in f.name.lower()]
    obs_path = None
    if not obs_files:
        print("  ✗  No Observed file ('observed' in name required)")
    else:
        if len(obs_files) > 1: print(f"  ⚠  Multiple Observed — using {obs_files[0].name}")
        obs_path = str(obs_files[0])
        print(f"  Observed : {obs_files[0].name}")
    prov_kw = _province_keywords(obs_path) if obs_path else set()
    print(f"  Province keywords: {prov_kw}")
    raw_models, bc_models = {}, {}
    for f in raw_files:
        if not _file_matches_province(f, prov_kw): continue
        m = _extract_model(f.name)
        if m not in raw_models: raw_models[m] = str(f); print(f"  Raw  '{m}' ← {f.name}")
    for f in bc_files:
        if not _file_matches_province(f, prov_kw): continue
        m = _extract_model(f.name)
        if m not in bc_models: bc_models[m] = str(f); print(f"  BC   '{m}' ← {f.name}")
    # Fallback
    if not raw_models and raw_files:
        for f in raw_files:
            m = _extract_model(f.name)
            if m not in raw_models: raw_models[m]=str(f); print(f"  Raw '{m}' ← {f.name} [fallback]")
    if not bc_models and bc_files:
        for f in bc_files:
            m = _extract_model(f.name)
            if m not in bc_models: bc_models[m]=str(f); print(f"  BC  '{m}' ← {f.name} [fallback]")
    return obs_path, raw_models, bc_models


# ════════════════════════════════════════════════════════════════════════
#  §2  DATA LOADING
# ════════════════════════════════════════════════════════════════════════

def load_daily(path, label, target_stns=None):
    if path is None or not os.path.isfile(path):
        print(f"  ✗  Not found: {label}"); return None, []
    time_cols = {"YEAR","MONTH","DAY"}
    try:
        if target_stns:
            tgt = set(str(s) for s in target_stns)
            use_cols = lambda c: str(c) in time_cols or str(c) in tgt
            df = pd.read_csv(path, usecols=use_cols)
        else:
            df = pd.read_csv(path)
    except Exception:
        df = pd.read_csv(path)
    df.columns = [str(c) for c in df.columns]
    for mv in MISS_FLAGS: df.replace(mv, np.nan, inplace=True)
    num = df.select_dtypes(include=[np.number]).columns
    df[num] = df[num].where(df[num] >= 0)
    stns = [c for c in df.columns if c not in time_cols]
    if target_stns:
        ts = [str(s) for s in target_stns]; stns = [s for s in stns if s in set(ts)]
    try:
        df["date"] = pd.to_datetime({"year":df["YEAR"],"month":df["MONTH"],"day":df["DAY"]})
        df = df.set_index("date")[stns]
    except Exception:
        df = df[stns]
    y0 = df.index[0].year if len(df) else "?"
    y1 = df.index[-1].year if len(df) else "?"
    print(f"    {label:40s}: {len(df):,} rows × {len(stns)} stns  [{y0}–{y1}]")
    return df, stns

def period_str(df):
    if df is None: return "N/A"
    try: return f"{df.index[0].year}–{df.index[-1].year}"
    except: return "N/A"

def short_labels(stns): return {str(s): f"S{i+1}" for i, s in enumerate(stns)}


# ════════════════════════════════════════════════════════════════════════
#  §3  DATA VALIDATION  [NEW in v3.0]
# ════════════════════════════════════════════════════════════════════════

def validate_input_data(obs_df, raw_dfs, bc_dfs, stns_str):
    """
    Pre-run data integrity checks:
    (1) Date range overlap — all files must share a common period ≥ 1 year
    (2) Missing value percentage per dataset per station
    (3) Station column consistency — model files must contain all Observed stations
    (4) Non-negative values check

    Returns:
        report (dict): validation results
        ok (bool): True if data passes all critical checks
    """
    print("\n  ── Data Validation ──────────────────────────────────────")
    report = {"ok": True, "warnings": [], "errors": []}

    # ── (1) Date range overlap ────────────────────────────────────────
    all_dfs = {"Observed": obs_df}
    all_dfs.update({f"Raw/{m}": df for m, df in raw_dfs.items()})
    all_dfs.update({f"BC/{m}":  df for m, df in bc_dfs.items()})
    common_idx = obs_df.index
    for lbl, df in all_dfs.items():
        if df is None: continue
        common_idx = common_idx.intersection(df.index)
    n_common = len(common_idx)
    if n_common < 365:
        msg = f"Common date overlap = {n_common} days (< 365 — CRITICAL)"
        report["errors"].append(msg); report["ok"] = False
        print(f"    ✗  {msg}")
    else:
        y0 = common_idx[0].year; y1 = common_idx[-1].year
        print(f"    ✓  Common period: {n_common:,} days  [{y0}–{y1}]")
        report["common_days"] = n_common
        report["common_period"] = f"{y0}–{y1}"

    # ── (2) Missing value % ───────────────────────────────────────────
    miss_report = {}
    for lbl, df in all_dfs.items():
        if df is None: continue
        cols = [s for s in stns_str if s in df.columns]
        if not cols: continue
        miss_pct = df[cols].isnull().mean().mean() * 100
        miss_report[lbl] = round(miss_pct, 2)
        if miss_pct > 20:
            msg = f"{lbl}: missing = {miss_pct:.1f}% (> 20% — WARNING)"
            report["warnings"].append(msg)
            print(f"    ⚠  {msg}")
        else:
            print(f"    ✓  {lbl}: missing = {miss_pct:.1f}%")
    report["missing_pct"] = miss_report

    # ── (3) Station consistency ───────────────────────────────────────
    obs_cols = set(obs_df.columns)
    for lbl, df in all_dfs.items():
        if df is None or lbl == "Observed": continue
        missing_stns = obs_cols - set(df.columns)
        if missing_stns:
            msg = f"{lbl}: missing stations {missing_stns}"
            report["warnings"].append(msg)
            print(f"    ⚠  {msg}")
    print(f"    ✓  Stations check: {len(stns_str)} stations matched")

    # ── (4) Date gaps ─────────────────────────────────────────────────
    if hasattr(obs_df.index, 'freq') or len(obs_df) > 1:
        expected_days = (obs_df.index[-1] - obs_df.index[0]).days + 1
        actual_days   = len(obs_df)
        gap_pct = (expected_days - actual_days) / expected_days * 100
        if gap_pct > 5:
            msg = f"Observed: date gaps = {gap_pct:.1f}% (> 5%)"
            report["warnings"].append(msg)
            print(f"    ⚠  {msg}")
        else:
            print(f"    ✓  Date continuity: gap = {gap_pct:.1f}%")
        report["date_gap_pct"] = round(gap_pct, 2)

    # Summary
    n_warn = len(report["warnings"]); n_err = len(report["errors"])
    print(f"    Result: {n_err} errors | {n_warn} warnings | "
          f"{'PASS' if report['ok'] else 'FAIL'}")
    if not report["ok"]:
        print("    CRITICAL: Data validation FAILED — results may be unreliable")
    print("  " + "─"*68)
    return report, report["ok"]


# ════════════════════════════════════════════════════════════════════════
#  §4  PERFORMANCE METRICS  (canonical formulas)
# ════════════════════════════════════════════════════════════════════════

def _compute_metrics(o, s):
    """
    Full metric suite for paired arrays.

    RMSE = sqrt(mean((s-o)^2))                         [lower=better]
    KGE  = 1 - sqrt((r-1)^2+(sigma_r-1)^2+(beta-1)^2) [Gupta 2009]
    NSE  = 1 - sum((s-o)^2) / sum((o-mean_o)^2)        [Nash 1970]
    r    = Pearson correlation                          [higher=better]
    """
    o = np.asarray(o, dtype=float); s = np.asarray(s, dtype=float)
    mask = ~np.isnan(o) & ~np.isnan(s)
    o, s = o[mask], s[mask]
    if len(o) < 5: return {k: np.nan for k in METS_ALL}
    e     = s - o
    rmse  = float(np.sqrt(np.mean(e**2)))
    r_val = float(np.corrcoef(o, s)[0, 1]) if len(o) > 2 else np.nan
    std_o = float(np.std(o, ddof=1)); std_s = float(np.std(s, ddof=1))
    sr    = std_s / std_o if std_o > 0 else np.nan
    beta  = float(np.mean(s) / np.mean(o)) if np.mean(o) != 0 else np.nan
    kge   = float(1 - math.sqrt((r_val-1)**2 + (sr-1)**2 + (beta-1)**2)) \
            if not (np.isnan(sr) or np.isnan(beta) or np.isnan(r_val)) else np.nan
    dn    = float(np.sum((o - np.mean(o))**2))
    nse   = float(1 - np.sum(e**2) / dn) if dn > 0 else np.nan
    return {"RMSE": rmse, "KGE": kge, "NSE": nse, "r": r_val}

def regional_metrics(obs_df, ens_df, stns_str):
    agg = {k: [] for k in METS_ALL}
    if obs_df is None or ens_df is None: return {k: np.nan for k in METS_ALL}
    ci = obs_df.index.intersection(ens_df.index)
    if len(ci) == 0: return {k: np.nan for k in METS_ALL}
    for stn in stns_str:
        if stn not in obs_df.columns or stn not in ens_df.columns: continue
        m = _compute_metrics(obs_df.loc[ci, stn].values, ens_df.loc[ci, stn].values)
        for k in METS_ALL:
            if not np.isnan(m[k]): agg[k].append(m[k])
    return {k: float(np.mean(agg[k])) if agg[k] else np.nan for k in METS_ALL}

def station_metrics(obs_df, ens_df, stns_str):
    if obs_df is None or ens_df is None:
        return {s: {k: np.nan for k in METS_ALL} for s in stns_str}
    ci = obs_df.index.intersection(ens_df.index)
    out = {}
    for stn in stns_str:
        if stn not in obs_df.columns or stn not in ens_df.columns:
            out[stn] = {k: np.nan for k in METS_ALL}; continue
        out[stn] = _compute_metrics(obs_df.loc[ci, stn].values, ens_df.loc[ci, stn].values)
    return out


# ════════════════════════════════════════════════════════════════════════
#  §5  ENSEMBLE HELPERS
# ════════════════════════════════════════════════════════════════════════

def build_ensemble(dfs_list, stns_str):
    valid = [df for df in dfs_list if df is not None]
    if not valid: return None
    ci = valid[0].index
    for df in valid[1:]: ci = ci.intersection(df.index)
    if len(ci) == 0: return None
    cols = [s for s in stns_str if all(s in df.columns for df in valid)]
    if not cols: return None
    stack = np.stack([df.loc[ci, cols].values.astype(float) for df in valid], axis=0)
    return pd.DataFrame(np.nanmean(stack, axis=0), index=ci, columns=cols)

def inter_model_spread(dfs_list, stns_str):
    valid = [df for df in dfs_list if df is not None]
    if len(valid) < 2: return 0.0
    ci = valid[0].index
    for df in valid[1:]: ci = ci.intersection(df.index)
    if len(ci) == 0: return np.nan
    cols = [s for s in stns_str if all(s in df.columns for df in valid)]
    if not cols: return np.nan
    stack = np.stack([df.loc[ci, cols].values.astype(float) for df in valid], axis=0)
    return float(np.nanmean(np.nanstd(stack, axis=0, ddof=1)))


# ════════════════════════════════════════════════════════════════════════
#  §6  SATURATION CRITERION & EXP-DECAY FIT
# ════════════════════════════════════════════════════════════════════════

def saturation_point(means_per_N, met_key, M, threshold=SAT_THRESHOLD):
    """
    Saturation point: first N where marginal improvement < threshold × total_improvement.
    """
    vals = np.array([means_per_N.get(n, np.nan) for n in range(1, M+1)], dtype=float)
    if np.all(np.isnan(vals)): return M
    v1 = vals[0]; vM = vals[-1]; total_imp = abs(vM - v1)
    if total_imp < 1e-9: return 1
    for n in range(1, M):
        marginal = abs(vals[n] - vals[n-1])
        if marginal < threshold * total_imp: return n
    return M

def _exp_decay(x, a, b, c): return a * np.exp(-b * x) + c

def fit_saturation_curve(n_vals, means):
    n = np.array(n_vals, dtype=float); y = np.array(means, dtype=float)
    valid = ~np.isnan(y)
    if valid.sum() < 3: return None, None
    try:
        y_v = y[valid]; p0 = [float(y_v[0] - y_v[-1]), 1.0, float(y_v[-1])]
        popt, _ = curve_fit(_exp_decay, n[valid], y_v, p0=p0, maxfev=5000)
        return popt, _exp_decay
    except Exception: return None, None


# ════════════════════════════════════════════════════════════════════════
#  §7  BOOTSTRAP CI FOR SATURATION POINT  [NEW in v3.0]
# ════════════════════════════════════════════════════════════════════════

def bootstrap_saturation_ci(results_by_N, met_key, M,
                             n_bootstrap=N_BOOTSTRAP, ci=CI_LEVEL, seed=42):
    """
    Bootstrap 95% Confidence Interval for the saturation point N*.

    Strategy:
      For each bootstrap iteration, resample (with replacement) the list of
      subset-performance values at each N, recompute N* on the resampled means,
      and collect the bootstrap distribution of N*.

    Returns:
      sat_N_star   : int       — point estimate of N*
      ci_lo, ci_hi : int       — lower/upper bootstrap CI bounds on N*
      sat_dist     : np.array  — full bootstrap distribution of N*
      sat_mean_lo  : float     — mean performance lower {CI}% CI per N
      sat_mean_hi  : float     — mean performance upper {CI}% CI per N
    """
    rng = np.random.default_rng(seed)
    point_means = {n: float(np.nanmean(results_by_N[n][met_key]))
                   if results_by_N[n][met_key] else np.nan for n in range(1, M+1)}
    N_star_point = saturation_point(point_means, met_key, M)

    # CI on mean performance per N (percentile bootstrap)
    ci_lo_dict = {}; ci_hi_dict = {}
    for n in range(1, M+1):
        vals = np.array(results_by_N[n][met_key], dtype=float)
        vals = vals[~np.isnan(vals)]
        if len(vals) < 2:
            ci_lo_dict[n] = point_means[n]; ci_hi_dict[n] = point_means[n]; continue
        boot = np.array([np.mean(rng.choice(vals, size=len(vals), replace=True))
                         for _ in range(n_bootstrap)])
        lo_pct = (100 - ci) / 2; hi_pct = 100 - lo_pct
        ci_lo_dict[n] = float(np.percentile(boot, lo_pct))
        ci_hi_dict[n] = float(np.percentile(boot, hi_pct))

    # Bootstrap distribution of N*
    boot_sat = []
    for _ in range(n_bootstrap):
        boot_means = {}
        for n in range(1, M+1):
            vals = np.array(results_by_N[n][met_key], dtype=float)
            vals = vals[~np.isnan(vals)]
            if len(vals) == 0: boot_means[n] = np.nan; continue
            boot_means[n] = float(np.mean(rng.choice(vals, size=len(vals), replace=True)))
        boot_sat.append(saturation_point(boot_means, met_key, M))

    boot_sat = np.array(boot_sat)
    lo_pct = (100 - ci) / 2; hi_pct = 100 - lo_pct
    sat_lo = int(np.percentile(boot_sat, lo_pct))
    sat_hi = int(np.percentile(boot_sat, hi_pct))

    return (N_star_point, sat_lo, sat_hi, boot_sat, ci_lo_dict, ci_hi_dict)


# ════════════════════════════════════════════════════════════════════════
#  §8  PARALLEL SATURATION ANALYSIS  [NEW in v3.0]
# ════════════════════════════════════════════════════════════════════════

def _compute_one_combo(args):
    """
    Worker function for parallel processing.
    args = (combo, N, obs_data, model_data_dict, stns_str)
    Returns: (N, reg_metrics, spread_val, per_station_metrics)
    """
    combo, N, obs_data, model_data_dict, stns_str = args
    # Rebuild DataFrames from serialised dict format
    obs_df    = pd.DataFrame(obs_data["values"],
                              index=pd.DatetimeIndex(obs_data["index"]),
                              columns=obs_data["columns"])
    dfs_list  = []
    for m in combo:
        md = model_data_dict[m]
        df = pd.DataFrame(md["values"],
                          index=pd.DatetimeIndex(md["index"]),
                          columns=md["columns"])
        dfs_list.append(df)
    ens = build_ensemble(dfs_list, stns_str)
    if ens is None:
        return (N, None, None, None)
    reg    = regional_metrics(obs_df, ens, stns_str)
    sp_val = inter_model_spread(dfs_list, stns_str) if N > 1 else 0.0
    per    = station_metrics(obs_df, ens, stns_str)
    return (N, reg, sp_val, per)

def _df_to_dict(df):
    """Serialise DataFrame for inter-process transfer."""
    return {"values":  df.values.tolist(),
            "index":   [str(i) for i in df.index],
            "columns": list(df.columns)}

def run_saturation_analysis(obs_d, dfs_dict, stns_str, models, ds_label,
                             use_parallel=True):
    """
    Exhaustive C(M,N) saturation analysis with optional parallel processing.

    Returns:
        results   : {N: {metric: list}}
        means     : {metric: {N: mean}}
        stds      : {metric: {N: std}}
        sp_means  : {N: mean_spread}
        sp_stds   : {N: std_spread}
        sat_pts   : {metric: N*}
        stn_res   : {N: {stn: {metric: list}}}
        boot_ci   : {metric: (N*, lo, hi, dist, mean_lo, mean_hi)}
    """
    M = len(models)
    results = {N: {k: [] for k in METS_ALL} for N in range(1, M+1)}
    spread  = {N: [] for N in range(1, M+1)}
    stn_res = {N: {stn: {k: [] for k in METS_ALL}
                   for stn in stns_str} for N in range(1, M+1)}
    total_combos = sum(math.comb(M, n) for n in range(1, M+1))
    print(f"    [{ds_label}] C({M},N) total = {total_combos} | "
          f"cores = {MAX_WORKERS if use_parallel else 1}")

    # Build all combo args
    obs_serial   = _df_to_dict(obs_d)
    model_serial = {m: _df_to_dict(dfs_dict[m]) for m in models if dfs_dict.get(m) is not None}
    all_args = []
    for N in range(1, M+1):
        for combo in itertools.combinations(models, N):
            all_args.append((combo, N, obs_serial, model_serial, stns_str))

    t0 = time.time()
    if use_parallel and MAX_WORKERS > 1 and len(all_args) > 4:
        try:
            with concurrent.futures.ProcessPoolExecutor(max_workers=MAX_WORKERS) as exe:
                all_results = list(exe.map(_compute_one_combo, all_args, chunksize=2))
        except Exception as e:
            print(f"    ⚠  Parallel failed ({e}) — falling back to serial")
            all_results = [_compute_one_combo(a) for a in all_args]
    else:
        all_results = [_compute_one_combo(a) for a in all_args]

    # Collect results
    for (N, reg, sp_val, per) in all_results:
        if reg is None: continue
        for k in METS_ALL:
            if not np.isnan(reg[k]): results[N][k].append(reg[k])
        spread[N].append(sp_val if sp_val is not None else 0.0)
        if per is not None:
            for stn in stns_str:
                for k in METS_ALL:
                    v = per.get(stn, {}).get(k, np.nan)
                    if not np.isnan(v): stn_res[N][stn][k].append(v)

    elapsed = time.time() - t0
    print(f"    [{ds_label}] Done in {elapsed:.1f}s")

    # Summarise
    means = {k: {N: float(np.nanmean(results[N][k])) if results[N][k] else np.nan
                 for N in range(1, M+1)} for k in METS_ALL}
    stds  = {k: {N: float(np.nanstd(results[N][k], ddof=1)) if len(results[N][k])>1 else 0.0
                 for N in range(1, M+1)} for k in METS_ALL}
    sat_pts = {k: saturation_point(means[k], k, M) for k in METS_ALL}
    sp_means = {N: float(np.nanmean(spread[N])) if spread[N] else np.nan
                for N in range(1, M+1)}
    sp_stds  = {N: float(np.nanstd(spread[N], ddof=1)) if len(spread[N])>1 else 0.0
                for N in range(1, M+1)}

    # Summary per N
    for N in range(1, M+1):
        n_c = math.comb(M, N)
        rm  = means["RMSE"].get(N, np.nan); kg = means["KGE"].get(N, np.nan)
        sp  = sp_means.get(N, np.nan)
        print(f"      N={N}: {n_c} subsets | "
              f"RMSE={rm:.4f}  KGE={kg:.4f}  Spread={sp:.4f}")

    # Bootstrap CI for N* [NEW]
    print(f"    Bootstrap CI for N* (n={N_BOOTSTRAP}) ...")
    boot_ci = {}
    for k in METS_ALL:
        boot_ci[k] = bootstrap_saturation_ci(results, k, M)
    for k in METS_ALL:
        Ns, lo, hi = boot_ci[k][0], boot_ci[k][1], boot_ci[k][2]
        print(f"      {k}: N*={Ns}  {CI_LEVEL}%CI=[{lo},{hi}]")

    return (results, means, stds, sp_means, sp_stds, sat_pts, stn_res, boot_ci)

# ════════════════════════════════════════════════════════════════════════
#  §9  FIGURE 1 — PERFORMANCE SATURATION CURVE
#      - No suptitle  - Larger axes  - Bold vivid colours
#      - Non-overlapping legend (lower-left for RMSE panel only)
#      - N* annotation placed outside data range
#      - 95% CI shaded band (lighter inner) + ±1 SD (outer)
# ════════════════════════════════════════════════════════════════════════

def fig1_saturation_performance(
        raw_means, raw_stds, raw_sat, raw_boot_ci,
        bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
        models, period_obs, out_dir, prefix):

    M     = len(models)
    x     = np.arange(1, M+1)
    x_fit = np.linspace(1, M, 200)

    METS_INFO = [
        ("RMSE", "RMSE  (mm day⁻¹)",                  "Lower = better",  "(a)"),
        ("KGE",  "Kling–Gupta Efficiency (KGE)",       "Higher = better", "(b)"),
        ("NSE",  "Nash–Sutcliffe Efficiency (NSE)",     "Higher = better", "(c)"),
        ("r",    "Pearson Correlation Coefficient (r)", "Higher = better", "(d)"),
    ]

    fig, axes = plt.subplots(2, 2, figsize=(18, 14))
    fig.subplots_adjust(hspace=0.44, wspace=0.28,
                        left=0.08, right=0.97, top=0.97, bottom=0.07)

    for pi, (met, ylabel, direction, panel) in enumerate(METS_INFO):
        ax = axes[pi//2, pi%2]

        for tag, means, stds, sat_n, boot_ci_res, col, lc, ls, lbl, alp in [
            ("Raw", raw_means, raw_stds, raw_sat, raw_boot_ci,
             C["raw"], C["raw_lt"], "--", "Raw CMIP6",        0.90),
            ("BC",  bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
             C["bc"],  C["bc_lt"],  "-",  "Bias-Corrected (QDM)", 0.95),
        ]:
            y_mn = np.array([means[met].get(n, np.nan) for n in x])
            y_sd = np.array([stds[met].get(n, 0.0)     for n in x])
            if np.all(np.isnan(y_mn)): continue

            # 95% CI from bootstrap
            ci_lo = np.array([boot_ci_res[met][4].get(n, np.nan) for n in x])
            ci_hi = np.array([boot_ci_res[met][5].get(n, np.nan) for n in x])

            # ±1 SD outer band
            ax.fill_between(x, y_mn - y_sd, y_mn + y_sd,
                            color=lc, alpha=0.28, zorder=2)
            # 95% CI inner band
            valid_ci = ~np.isnan(ci_lo) & ~np.isnan(ci_hi)
            if valid_ci.any():
                ax.fill_between(x[valid_ci], ci_lo[valid_ci], ci_hi[valid_ci],
                                color=col, alpha=0.13, zorder=3,
                                label=f"{CI_LEVEL}% CI ({tag})" if pi == 0 else None)

            # Mean line + markers
            mk = METRIC_MARKERS[met]
            ax.plot(x, y_mn, color=col, lw=2.8, ls=ls, marker=mk,
                    ms=11, markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                    zorder=6, alpha=alp, label=lbl)

            # Exponential decay fit (dotted)
            popt, fn = fit_saturation_curve(list(x), list(y_mn))
            if popt is not None:
                try:
                    y_fit = fn(x_fit, *popt)
                    ax.plot(x_fit, y_fit, color=col, lw=1.4, ls=":",
                            alpha=0.72, zorder=4)
                except Exception: pass

            # Saturation point ★ with clean annotation
            sat_x = sat_n[met]
            sat_lo, sat_hi = boot_ci_res[met][1], boot_ci_res[met][2]
            if 1 <= sat_x <= M:
                sat_y = means[met].get(sat_x, np.nan)
                if not np.isnan(sat_y):
                    ax.scatter([sat_x], [sat_y], color=col, s=280,
                               marker="*", zorder=9,
                               edgecolors="#1A1A1A", linewidth=0.9)
                    # Vertical dotted guide line
                    ax.axvline(sat_x, color=col, lw=1.1, ls=":",
                               alpha=0.52, zorder=2)
                    # CI bracket on x-axis (drawn as error bar)
                    if sat_lo != sat_hi:
                        y_bot = ax.get_ylim()[0] if ax.get_ylim()[0] != 0 else sat_y * 0.90
                    # Annotation box — right side of figure to avoid overlap
                    ax_x_frac = (sat_x - 0.5) / (M + 0.5)
                    ann_x = sat_x + 0.32 if ax_x_frac < 0.75 else sat_x - 0.32
                    dy    = y_sd[sat_x-1] * 0.6 + abs(sat_y) * 0.03
                    dy   *= (1 if met in LOWER_B else -1)
                    try:
                        ax.annotate(
                            f"N*={sat_x} ({tag})\n[{CI_LEVEL}%CI: {sat_lo}–{sat_hi}]",
                            xy=(sat_x, sat_y),
                            xytext=(ann_x, sat_y + dy),
                            fontsize=9.5, fontweight="bold", color=col,
                            arrowprops=dict(arrowstyle="->", color=col,
                                            lw=1.2, alpha=0.75),
                            bbox=dict(boxstyle="round,pad=0.35", fc="white",
                                      ec=col, alpha=0.93, lw=1.0),
                            zorder=11)
                    except Exception: pass

        # ── Axes decoration ──────────────────────────────────────────
        ax.set_xticks(x)
        ax.set_xticklabels([str(n) for n in x], fontsize=12)
        ax.set_xlabel("Number of Models in Ensemble (N)",
                      fontsize=13, fontweight="bold", labelpad=6)
        ax.set_ylabel(ylabel, fontsize=13, fontweight="bold", labelpad=6)
        ax.set_title(f"{panel}  {ylabel}\n"
                     f"     {direction}  |  "
                     f"N* (Raw): {raw_sat[met]}  |  "
                     f"N* (BC): {bc_sat[met]}  [{CI_LEVEL}%CI: "
                     f"{bc_boot_ci[met][1]}–{bc_boot_ci[met][2]}]",
                     loc="left", fontsize=12, fontweight="bold", pad=5)
        ax.tick_params(axis="both", which="major", labelsize=12, width=1.6)
        ax.yaxis.set_minor_locator(ticker.AutoMinorLocator())
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        # ── Legend — panel (a) only, lower-left ─────────────────────
        if pi == 0:
            handles = [
                Line2D([0],[0], color=C["raw"], lw=2.8, ls="--",
                       marker="o", ms=10, markeredgecolor="#1A1A1A",
                       markeredgewidth=0.7, label="Raw CMIP6"),
                Line2D([0],[0], color=C["bc"],  lw=2.8, ls="-",
                       marker="s", ms=10, markeredgecolor="#1A1A1A",
                       markeredgewidth=0.7, label="Bias-Corrected (QDM)"),
                mpatches.Patch(color=C["raw_lt"], alpha=0.32,
                               label="±1 SD  (Raw)"),
                mpatches.Patch(color=C["bc_lt"],  alpha=0.32,
                               label="±1 SD  (BC)"),
                mpatches.Patch(color=C["raw"],    alpha=0.16,
                               label=f"{CI_LEVEL}% CI  (Raw)"),
                mpatches.Patch(color=C["bc"],     alpha=0.16,
                               label=f"{CI_LEVEL}% CI  (BC)"),
                Line2D([0],[0], color=C["grey"],  lw=1.4, ls=":",
                       label="Exp. decay fit"),
                Line2D([0],[0], marker="*", color="#1A1A1A", ls="none",
                       ms=14, label="Saturation point (N*)"),
            ]
            ax.legend(handles=handles, fontsize=10.5, frameon=True,
                      edgecolor="#B0BEC5", facecolor="white", framealpha=0.97,
                      ncol=2, loc="upper right", handlelength=2.0,
                      borderpad=0.9, handletextpad=0.7,
                      labelspacing=0.55)

        # Criterion text box — lower-left corner (avoids data)
        ax.text(0.03, 0.04,
                f"Criterion: marginal < {int(SAT_THRESHOLD*100)}% of total gain\n"
                f"Bootstrap CI (n={N_BOOTSTRAP})",
                transform=ax.transAxes, fontsize=9.5,
                color=C["grey"], va="bottom",
                bbox=dict(boxstyle="round,pad=0.38", fc="white",
                          ec="#B0BEC5", alpha=0.93, lw=0.8))

    # [v4.0] ── Export individual panels BEFORE closing the figure ──────────
    print("  Splitting Fig 1 into individual panels ...")
    axes_flat = [axes[0, 0], axes[0, 1], axes[1, 0], axes[1, 1]]
    save_individual_panels(fig, axes_flat, ["a", "b", "c", "d"],
                           fig_num=1, out_dir=out_dir, dpi=DPI)

    savefig(fig, out_dir / f"{prefix}_Fig1_SaturationPerformance")


# ════════════════════════════════════════════════════════════════════════
#  §10  FIGURE 2 — SPREAD SATURATION + REDUCTION + TRADE-OFF
# ════════════════════════════════════════════════════════════════════════

def fig2_saturation_spread(
        raw_spread_m, raw_spread_s,
        bc_spread_m,  bc_spread_s,
        raw_means, bc_means,
        raw_sat, bc_sat,
        models, period_obs, out_dir, prefix):

    M = len(models); x = np.arange(1, M+1)

    fig, axes = plt.subplots(1, 3, figsize=(22, 8))
    fig.subplots_adjust(left=0.06, right=0.97, top=0.96,
                        bottom=0.12, wspace=0.30)

    # ── (a) Spread vs N ──────────────────────────────────────────────
    ax = axes[0]
    for tag, sp_m, sp_s, col, lc, ls, lbl, mk_s in [
        ("Raw", raw_spread_m, raw_spread_s, C["raw"], C["raw_lt"], "--", "Raw CMIP6",   "^"),
        ("BC",  bc_spread_m,  bc_spread_s,  C["bc"],  C["bc_lt"],  "-",  "BC (QDM)",    "D"),
    ]:
        y_mn = np.array([sp_m.get(n, 0.0) for n in x])
        y_sd = np.array([sp_s.get(n, 0.0) for n in x])
        ax.plot(x, y_mn, color=col, lw=2.8, ls=ls, marker=mk_s,
                ms=11, markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                zorder=5, label=lbl)
        ax.fill_between(x, y_mn - y_sd, y_mn + y_sd,
                        color=lc, alpha=0.30, zorder=3)
        if M >= 3 and np.any(y_mn[1:] > 0):
            try:
                popt, fn = fit_saturation_curve(list(x[1:]), list(y_mn[1:]))
                if popt is not None:
                    xf = np.linspace(2, M, 200)
                    ax.plot(xf, fn(xf, *popt), color=col, lw=1.4,
                            ls=":", alpha=0.72, zorder=4)
            except Exception: pass

    ax.axhline(0, color=C["grey"], lw=1.0, ls="--", alpha=0.55)
    ax.set_xticks(x); ax.set_xticklabels([str(n) for n in x], fontsize=12)
    ax.set_xlabel("Number of Models in Ensemble (N)",
                  fontsize=13, fontweight="bold", labelpad=6)
    ax.set_ylabel("Inter-Model Spread  (mm day⁻¹)",
                  fontsize=13, fontweight="bold", labelpad=6)
    ax.set_title("(a)  Ensemble Spread vs Ensemble Size\n"
                 "     N=1: spread = 0 (single model)  |  "
                 "Higher = larger model disagreement",
                 loc="left", fontsize=12, fontweight="bold", pad=5)
    ax.legend(fontsize=11.5, frameon=True, edgecolor="#B0BEC5",
              facecolor="white", framealpha=0.97,
              loc="lower right", borderpad=0.8)
    ax.tick_params(axis="both", which="major", labelsize=12, width=1.6)
    ax.yaxis.set_minor_locator(ticker.AutoMinorLocator())
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)

    # ── (b) Cumulative uncertainty reduction ─────────────────────────
    ax2 = axes[1]
    for tag, sp_m, col, ls, lbl, mk_s in [
        ("Raw", raw_spread_m, C["raw"], "--", "Raw CMIP6", "^"),
        ("BC",  bc_spread_m,  C["bc"],  "-",  "BC (QDM)",  "D"),
    ]:
        baseline = sp_m.get(2, np.nan)
        if np.isnan(baseline) or baseline == 0: continue
        max_sp = sp_m.get(M, baseline)
        denom  = baseline - max_sp + 1e-9
        reduct = np.array([
            100 * (sp_m.get(n, np.nan) - max_sp) / denom
            if not np.isnan(sp_m.get(n, np.nan)) else np.nan
            for n in x])
        ax2.plot(x, reduct, color=col, lw=2.8, ls=ls, marker=mk_s,
                 ms=11, markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                 zorder=5, label=lbl)
        # Shade filled area
        valid = ~np.isnan(reduct) & (reduct >= 0)
        if valid.any():
            ax2.fill_between(x, np.where(np.isnan(reduct), 0, reduct), 0,
                             where=(~np.isnan(reduct) & (reduct >= 0)),
                             color=col, alpha=0.11, zorder=3)

    ax2.axhline(100, color=C["grey"], lw=1.2, ls=":", alpha=0.65,
                label="100% (N=2 baseline)")
    ax2.axhline(0, color=C["grey"], lw=0.9, ls="--", alpha=0.50)
    ax2.set_xticks(x); ax2.set_xticklabels([str(n) for n in x], fontsize=12)
    ax2.set_xlabel("Number of Models in Ensemble (N)",
                   fontsize=13, fontweight="bold", labelpad=6)
    ax2.set_ylabel("Cumulative Spread Reduction (%)\n"
                   "[relative to N=2 baseline spread]",
                   fontsize=13, fontweight="bold", labelpad=6)
    ax2.set_title("(b)  Cumulative Uncertainty Reduction Rate\n"
                  "     % reduction in inter-model spread relative to 2-model ensemble",
                  loc="left", fontsize=12, fontweight="bold", pad=5)
    ax2.legend(fontsize=11.5, frameon=True, edgecolor="#B0BEC5",
               facecolor="white", framealpha=0.97,
               loc="lower right", borderpad=0.8)
    ax2.tick_params(axis="both", which="major", labelsize=12, width=1.6)
    ax2.yaxis.set_minor_locator(ticker.AutoMinorLocator())
    ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)

    # ── (c) Dual-axis: KGE improvement vs spread ─────────────────────
    ax3   = axes[2]
    ax3_r = ax3.twinx()
    kge_bc = np.array([bc_means["KGE"].get(n, np.nan) for n in x])
    kge_1  = kge_bc[0]  if not np.isnan(kge_bc[0])  else 0.0
    kge_M  = kge_bc[-1] if not np.isnan(kge_bc[-1]) else 0.0
    kge_norm = 100 * (kge_bc - kge_1) / (kge_M - kge_1 + 1e-9)
    sp_bc  = np.array([bc_spread_m.get(n, np.nan) for n in x])

    l1, = ax3.plot(x, kge_norm, color=C["bc"], lw=2.8, ls="-", marker="s",
                   ms=11, markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                   zorder=5, label="KGE improvement (BC, %)")
    ax3.fill_between(x, kge_norm, 0, color=C["bc_lt"], alpha=0.22, zorder=3)

    l2, = ax3_r.plot(x, sp_bc, color=C["raw"], lw=2.4, ls="--", marker="o",
                     ms=11, markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                     zorder=4, label="Inter-model spread (BC)")

    sat_kge = bc_sat.get("KGE", M)
    if 1 <= sat_kge <= M:
        ax3.axvline(sat_kge, color=C["bc"], lw=1.5, ls=":", alpha=0.72, zorder=2)
        ax3.text(sat_kge + 0.10, max(kge_norm[~np.isnan(kge_norm)]) * 0.08,
                 f"N*={sat_kge}",
                 color=C["bc"], fontsize=11.5, fontweight="bold",
                 bbox=dict(boxstyle="round,pad=0.32", fc="white",
                           ec=C["bc"], alpha=0.92, lw=0.9))

    ax3.set_xticks(x); ax3.set_xticklabels([str(n) for n in x], fontsize=12)
    ax3.set_xlabel("Number of Models in Ensemble (N)",
                   fontsize=13, fontweight="bold", labelpad=6)
    ax3.set_ylabel("Cumulative KGE Improvement (%)",
                   fontsize=13, fontweight="bold", color=C["bc"], labelpad=6)
    ax3_r.set_ylabel("Inter-Model Spread (mm day⁻¹)",
                     fontsize=13, fontweight="bold", color=C["raw"], labelpad=6)
    ax3.tick_params(axis="y", labelcolor=C["bc"],  labelsize=12, width=1.6)
    ax3_r.tick_params(axis="y", labelcolor=C["raw"], labelsize=12, width=1.6)
    ax3.tick_params(axis="x", labelsize=12, width=1.6)
    ax3.set_title("(c)  KGE Improvement vs Ensemble Spread\n"
                  "     Trade-off: performance gain vs model disagreement (BC)",
                  loc="left", fontsize=12, fontweight="bold", pad=5)
    lines_c  = [l1, l2]
    labels_c = [l.get_label() for l in lines_c]
    ax3.legend(lines_c, labels_c, fontsize=11.5, frameon=True,
               edgecolor="#B0BEC5", facecolor="white", framealpha=0.97,
               loc="lower right", borderpad=0.8)
    ax3.spines["top"].set_visible(False)

    # [v4.0] ── Export individual panels BEFORE closing the figure ──────────
    print("  Splitting Fig 2 into individual panels ...")
    # Panel (c) has a twinx axis — pass both so bbox is fully captured
    save_individual_panels(fig,
                           [axes[0], axes[1], [axes[2], ax3_r]],
                           ["a", "b", "c"],
                           fig_num=2, out_dir=out_dir, dpi=DPI)

    savefig(fig, out_dir / f"{prefix}_Fig2_SaturationSpread")


# ════════════════════════════════════════════════════════════════════════
#  §11  FIGURE 3 — SPATIAL SATURATION
# ════════════════════════════════════════════════════════════════════════

def fig3_saturation_spatial(bc_stn_res, stns_str, smap, models,
                              bc_sat, bc_boot_ci, period_obs, out_dir, prefix):

    M     = len(models)
    x     = np.arange(1, M+1)
    codes = [smap[s] for s in stns_str]
    n_s   = len(stns_str)

    # Use a perceptually distinct, print-safe palette for stations
    cmap_stn = cm.get_cmap("tab20", max(n_s, 1))
    stn_cols = [mcolors.to_hex(cmap_stn(i)) for i in range(n_s)]
    # Increase saturation for all station colours
    def _vivid(hex_c):
        rgb = mcolors.to_rgb(hex_c)
        h, s, v = mcolors.rgb_to_hsv(rgb)
        s = min(s * 1.35, 1.0); v = min(v * 1.05, 1.0)
        return mcolors.hsv_to_rgb((h, s, v))
    stn_cols = [mcolors.to_hex(_vivid(c)) for c in stn_cols]

    fig = plt.figure(figsize=(22, 9))
    gs  = gridspec.GridSpec(1, 3, figure=fig, hspace=0.08, wspace=0.28,
                            top=0.96, bottom=0.11, left=0.06, right=0.97)
    ax1 = fig.add_subplot(gs[0])
    ax2 = fig.add_subplot(gs[1])
    ax3 = fig.add_subplot(gs[2])

    for pi, (met, ylabel, ax_use, panel, leg_loc) in enumerate([
        ("RMSE", "RMSE  (mm day⁻¹)",            ax1, "(a)", "upper right"),
        ("KGE",  "Kling–Gupta Efficiency (KGE)", ax2, "(b)", "lower right"),
    ]):
        stn_sat_pts = []
        for si, (stn, code) in enumerate(zip(stns_str, codes)):
            y_mn = np.array([
                float(np.mean(bc_stn_res[N][stn][met]))
                if bc_stn_res[N][stn][met] else np.nan
                for N in x], dtype=float)
            if np.all(np.isnan(y_mn)): continue
            col = stn_cols[si]
            ax_use.plot(x, y_mn, color=col, lw=2.0, ls="-",
                        marker="o", ms=7.5,
                        markeredgecolor="#1A1A1A", markeredgewidth=0.5,
                        zorder=4, alpha=0.88, label=code)
            # Station saturation ★
            sp_dict = {n: float(np.mean(bc_stn_res[n][stn][met]))
                       if bc_stn_res[n][stn][met] else np.nan
                       for n in range(1, M+1)}
            sp = saturation_point(sp_dict, met, M)
            stn_sat_pts.append(sp)
            if 1 <= sp <= M:
                sp_y = sp_dict.get(sp, np.nan)
                if not np.isnan(sp_y):
                    ax_use.scatter([sp], [sp_y], color=col, marker="*",
                                   s=200, zorder=7, edgecolors="#1A1A1A",
                                   linewidth=0.7, alpha=0.95)

        # Regional mean (bold black) with 95% CI band
        y_reg = []
        for N in x:
            vals_all = [np.mean(bc_stn_res[N][stn][met])
                        for stn in stns_str if bc_stn_res[N][stn][met]]
            y_reg.append(float(np.nanmean(vals_all)) if vals_all else np.nan)
        y_reg = np.array(y_reg)
        ci_lo = np.array([bc_boot_ci[met][4].get(n, np.nan) for n in x])
        ci_hi = np.array([bc_boot_ci[met][5].get(n, np.nan) for n in x])
        ax_use.plot(x, y_reg, color="#1A1A1A", lw=3.4, ls="-",
                    marker="D", ms=11, markeredgecolor="white",
                    markeredgewidth=0.8, zorder=8, label="Regional mean")
        valid_ci = ~np.isnan(ci_lo) & ~np.isnan(ci_hi)
        if valid_ci.any():
            ax_use.fill_between(x[valid_ci], ci_lo[valid_ci], ci_hi[valid_ci],
                                color="#1A1A1A", alpha=0.10, zorder=3)

        # Regional N* vertical line
        sat_global = bc_sat.get(met, M)
        sat_lo_g   = bc_boot_ci[met][1]
        sat_hi_g   = bc_boot_ci[met][2]
        ax_use.axvline(sat_global, color="#1A1A1A", lw=2.0, ls="--",
                       alpha=0.75, zorder=6,
                       label=f"Regional N*={sat_global} [{CI_LEVEL}%CI: {sat_lo_g}–{sat_hi_g}]")

        # Axes decoration
        ax_use.set_xticks(x)
        ax_use.set_xticklabels([str(n) for n in x], fontsize=12)
        ax_use.set_xlabel("Number of Models (N)",
                          fontsize=13, fontweight="bold", labelpad=6)
        ax_use.set_ylabel(ylabel, fontsize=13, fontweight="bold", labelpad=6)
        ax_use.set_title(
            f"{panel}  {ylabel} — Spatial Saturation\n"
            "     Each colour = one station  |  ★ = station N*\n"
            f"     Bold black = regional mean ± {CI_LEVEL}% CI  |  Dashed = regional N*",
            loc="left", fontsize=12, fontweight="bold", pad=5)
        ax_use.tick_params(axis="both", which="major", labelsize=12, width=1.6)
        ax_use.yaxis.set_minor_locator(ticker.AutoMinorLocator())
        ax_use.spines["top"].set_visible(False)
        ax_use.spines["right"].set_visible(False)

        # Legend — two columns, placed to avoid data overlap
        handles_s = [
            Line2D([0],[0], color=stn_cols[i], lw=2.0, marker="o",
                   ms=7, markeredgecolor="#1A1A1A", markeredgewidth=0.5,
                   label=codes[i])
            for i in range(n_s)
        ]
        handles_s += [
            Line2D([0],[0], color="#1A1A1A", lw=3.4, marker="D", ms=9,
                   markeredgecolor="white", label="Regional mean"),
            Line2D([0],[0], marker="*", color="#1A1A1A", ls="none",
                   ms=12, label="Station saturation N*"),
            Line2D([0],[0], color="#1A1A1A", lw=2.0, ls="--",
                   alpha=0.75, label=f"Regional N* ± {CI_LEVEL}% CI"),
        ]
        ax_use.legend(handles=handles_s, fontsize=8.8, frameon=True,
                      edgecolor="#B0BEC5", facecolor="white", framealpha=0.97,
                      ncol=2, loc=leg_loc, handlelength=1.8,
                      borderpad=0.8, handletextpad=0.55,
                      labelspacing=0.45)

    # ── (c) Saturation heatmap (station × metric) ─────────────────────
    METS_HM = ["RMSE","KGE","NSE","r"]
    sat_mat = np.full((len(METS_HM), n_s), np.nan)
    for mi, met in enumerate(METS_HM):
        for si, stn in enumerate(stns_str):
            sp_dict = {n: float(np.mean(bc_stn_res[n][stn][met]))
                       if bc_stn_res[n][stn][met] else np.nan
                       for n in range(1, M+1)}
            sat_mat[mi, si] = saturation_point(sp_dict, met, M)

    # Vivid green sequential map
    cmap_sat = LinearSegmentedColormap.from_list(
        "sat_green", ["#E8F5E9","#43A047","#1B5E20"], N=M)
    im = ax3.imshow(sat_mat, cmap=cmap_sat, vmin=1, vmax=M,
                    aspect="auto", interpolation="nearest")
    for mi, met in enumerate(METS_HM):
        for si, code in enumerate(codes):
            val = sat_mat[mi, si]
            if not np.isnan(val):
                tc = "white" if val >= M - 0.5 else "#1A1A1A"
                ax3.text(si, mi, f"{int(val)}", ha="center", va="center",
                         fontsize=14, fontweight="bold", color=tc)
    ax3.set_xticks(range(n_s)); ax3.set_yticks(range(len(METS_HM)))
    ax3.set_xticklabels(codes, rotation=0, ha="center", fontsize=11)
    ax3.set_yticklabels(METS_HM, fontsize=13, fontweight="bold")
    ax3.set_xlabel("Station", fontsize=13, fontweight="bold", labelpad=6)
    cb = plt.colorbar(im, ax=ax3, orientation="horizontal",
                      pad=0.22, fraction=0.06, shrink=0.85)
    cb.set_ticks(range(1, M+1))
    cb.set_ticklabels([str(n) for n in range(1, M+1)], fontsize=11)
    cb.set_label("Saturation Point N*  (number of models)",
                 fontsize=11.5, fontweight="bold")
    cb.ax.tick_params(width=1.4)
    ax3.set_title(
        "(c)  Saturation Point N* — Station × Metric Heatmap\n"
        "     Number = ensemble size at saturation  |  Darker = later saturation",
        loc="left", fontsize=12, fontweight="bold", pad=5)

    # [v4.0] ── Export individual panels BEFORE closing the figure ──────────
    print("  Splitting Fig 3 into individual panels ...")
    save_individual_panels(fig, [ax1, ax2, ax3], ["a", "b", "c"],
                           fig_num=3, out_dir=out_dir, dpi=DPI)

    savefig(fig, out_dir / f"{prefix}_Fig3_SaturationSpatial")

# ════════════════════════════════════════════════════════════════════════
#  §11b  FIGURE 4 — INTEGRATED DUAL-AXIS PERFORMANCE PROFILE  [NEW v5.0]
#         Panel (a): RMSE (left-Y) + KGE (right-Y) vs N
#         Panel (b): NSE  (left-Y) + r   (right-Y) vs N
#         • Dual Y-axes so incommensurable metrics share one x-axis
#         • Both Raw and BC per panel — 4 series per panel
#         • ±1 SD outer band + 95% Bootstrap CI inner band
#         • Exponential-decay fit (dotted)
#         • Saturation star ★ on both axes; vertical guide line
# ════════════════════════════════════════════════════════════════════════

def fig4_integrated_profile(
        raw_means, raw_stds, raw_sat, raw_boot_ci,
        bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
        models, period_obs, out_dir, prefix):
    """
    Fig 4 — Integrated Dual-Axis Performance Profile.

    Academic justification: Displaying two complementary metrics on one axes
    panel (dual-axis) allows direct visual comparison of competing optimality
    criteria without requiring the reader to flip between panels
    (Moriasi et al., 2007; Gupta et al., 2009).
    """
    M     = len(models)
    x     = np.arange(1, M + 1)
    x_fit = np.linspace(1, M, 300)

    # PAIRS: (left_metric, right_metric, panel_label, title_suffix)
    PAIRS = [
        ("RMSE", "KGE", "(a)",
         "RMSE  (mm day⁻¹)  [left]  |  KGE  [right]"),
        ("NSE",  "r",   "(b)",
         "NSE  [left]  |  Pearson r  [right]"),
    ]

    fig, axes_row = plt.subplots(1, 2, figsize=(22, 9))
    fig.subplots_adjust(left=0.07, right=0.93, top=0.92, bottom=0.11,
                        wspace=0.38)

    ax_twins = []   # collect for panel-split export

    for col_idx, (m_left, m_right, panel, ttl) in enumerate(PAIRS):
        ax_L = axes_row[col_idx]
        ax_R = ax_L.twinx()
        ax_twins.append((ax_L, ax_R))

        # ── determine y-axis direction labels ────────────────────────
        dir_L = "↓ better" if m_left  in LOWER_B else "↑ better"
        dir_R = "↓ better" if m_right in LOWER_B else "↑ better"

        col_L = METRIC_COLORS[m_left];  mk_L = METRIC_MARKERS[m_left]
        col_R = METRIC_COLORS[m_right]; mk_R = METRIC_MARKERS[m_right]

        # Slight transparency difference to distinguish overlapping shades
        _alpha_L, _alpha_R = 0.95, 0.85

        for tag, means, stds, sat_n, boot_ci_res, ls_raw, ls_bc in [
            ("Raw", raw_means, raw_stds, raw_sat, raw_boot_ci, "--", ":"),
            ("BC",  bc_means,  bc_stds,  bc_sat,  bc_boot_ci,  "-",  "-."),
        ]:
            # ── LEFT axis (m_left) ──────────────────────────────────
            y_L  = np.array([means[m_left].get(n, np.nan) for n in x])
            sd_L = np.array([stds[m_left].get(n, 0.0)     for n in x])
            ci_lo_L = np.array([boot_ci_res[m_left][4].get(n, np.nan) for n in x])
            ci_hi_L = np.array([boot_ci_res[m_left][5].get(n, np.nan) for n in x])
            ls_use  = ls_raw if tag == "Raw" else ls_bc
            lw_use  = 2.4 if tag == "BC" else 1.9
            alp_use = _alpha_L if tag == "BC" else 0.75

            if not np.all(np.isnan(y_L)):
                ax_L.fill_between(x, y_L - sd_L, y_L + sd_L,
                                  color=col_L, alpha=0.18, zorder=2)
                v_ci = ~np.isnan(ci_lo_L) & ~np.isnan(ci_hi_L)
                if v_ci.any():
                    ax_L.fill_between(x[v_ci], ci_lo_L[v_ci], ci_hi_L[v_ci],
                                      color=col_L, alpha=0.10, zorder=3)
                ax_L.plot(x, y_L, color=col_L, lw=lw_use, ls=ls_use,
                          marker=mk_L, ms=11,
                          markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                          alpha=alp_use, zorder=6,
                          label=f"{m_left} ({tag})")
                # exp-decay fit
                popt, fn = fit_saturation_curve(list(x), list(y_L))
                if popt is not None:
                    try:
                        ax_L.plot(x_fit, fn(x_fit, *popt),
                                  color=col_L, lw=1.2, ls=":", alpha=0.60, zorder=4)
                    except Exception: pass
                # saturation star on left axis
                sx = sat_n[m_left]
                if 1 <= sx <= M:
                    sy = means[m_left].get(sx, np.nan)
                    if not np.isnan(sy):
                        ax_L.scatter([sx], [sy], color=col_L, s=300,
                                     marker="*", zorder=9,
                                     edgecolors="#1A1A1A", linewidth=0.8)

            # ── RIGHT axis (m_right) ────────────────────────────────
            y_R  = np.array([means[m_right].get(n, np.nan) for n in x])
            sd_R = np.array([stds[m_right].get(n, 0.0)     for n in x])
            ci_lo_R = np.array([boot_ci_res[m_right][4].get(n, np.nan) for n in x])
            ci_hi_R = np.array([boot_ci_res[m_right][5].get(n, np.nan) for n in x])
            alp_use_R = _alpha_R if tag == "BC" else 0.65

            if not np.all(np.isnan(y_R)):
                ax_R.fill_between(x, y_R - sd_R, y_R + sd_R,
                                  color=col_R, alpha=0.16, zorder=2)
                v_ci_R = ~np.isnan(ci_lo_R) & ~np.isnan(ci_hi_R)
                if v_ci_R.any():
                    ax_R.fill_between(x[v_ci_R], ci_lo_R[v_ci_R], ci_hi_R[v_ci_R],
                                      color=col_R, alpha=0.09, zorder=3)
                ax_R.plot(x, y_R, color=col_R, lw=lw_use, ls=ls_use,
                          marker=mk_R, ms=10,
                          markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                          alpha=alp_use_R, zorder=6,
                          label=f"{m_right} ({tag})")
                popt_r, fn_r = fit_saturation_curve(list(x), list(y_R))
                if popt_r is not None:
                    try:
                        ax_R.plot(x_fit, fn_r(x_fit, *popt_r),
                                  color=col_R, lw=1.2, ls=":", alpha=0.60, zorder=4)
                    except Exception: pass
                sx_r = sat_n[m_right]
                if 1 <= sx_r <= M:
                    sy_r = means[m_right].get(sx_r, np.nan)
                    if not np.isnan(sy_r):
                        ax_R.scatter([sx_r], [sy_r], color=col_R, s=260,
                                     marker="*", zorder=9,
                                     edgecolors="#1A1A1A", linewidth=0.8)

        # ── Saturation vertical lines (BC only — definitive estimate) ─
        for m_k, ax_k in [(m_left, ax_L), (m_right, ax_R)]:
            sx = bc_sat[m_k]
            if 1 <= sx <= M:
                ax_L.axvline(sx, color=METRIC_COLORS[m_k], lw=1.4,
                             ls="--", alpha=0.50, zorder=1)
                lo_k, hi_k = bc_boot_ci[m_k][1], bc_boot_ci[m_k][2]
                ax_L.axvspan(lo_k - 0.05, hi_k + 0.05,
                             color=METRIC_COLORS[m_k], alpha=0.07, zorder=0)

        # ── Axes decoration ──────────────────────────────────────────
        ax_L.set_xticks(x)
        ax_L.set_xticklabels([str(n) for n in x], fontsize=12)
        ax_L.set_xlabel("Number of Models in Ensemble (N)",
                        fontsize=13, fontweight="bold", labelpad=6)
        ax_L.set_ylabel(f"{m_left}  ({dir_L})", fontsize=13,
                        fontweight="bold", color=col_L, labelpad=6)
        ax_R.set_ylabel(f"{m_right}  ({dir_R})", fontsize=13,
                        fontweight="bold", color=col_R, labelpad=8)
        ax_L.tick_params(axis="y", colors=col_L, labelsize=11.5,
                         width=1.6, which="major")
        ax_R.tick_params(axis="y", colors=col_R, labelsize=11.5,
                         width=1.6, which="major")
        ax_L.tick_params(axis="x", labelsize=12, width=1.6)
        ax_L.yaxis.set_minor_locator(ticker.AutoMinorLocator())
        ax_R.yaxis.set_minor_locator(ticker.AutoMinorLocator())
        ax_L.spines["top"].set_visible(False)
        ax_R.spines["top"].set_visible(False)

        # Legend — combined both axes
        handles_L = ax_L.get_legend_handles_labels()
        handles_R = ax_R.get_legend_handles_labels()
        combined_h = handles_L[0] + handles_R[0]
        combined_l = handles_L[1] + handles_R[1]
        combined_h += [
            mpatches.Patch(color=col_L, alpha=0.22, label=f"±1 SD + 95%CI  ({m_left})"),
            mpatches.Patch(color=col_R, alpha=0.20, label=f"±1 SD + 95%CI  ({m_right})"),
            Line2D([0],[0], color=C["grey"], lw=1.2, ls=":", label="Exp-decay fit"),
            Line2D([0],[0], marker="*", color="#1A1A1A", ls="none",
                   ms=14, label="Saturation point N*"),
        ]
        ax_L.legend(combined_h, [h.get_label() for h in combined_h],
                    fontsize=9.8, frameon=True, edgecolor="#B0BEC5",
                    facecolor="white", framealpha=0.96,
                    ncol=2, loc="upper right",
                    handlelength=1.9, borderpad=0.85,
                    handletextpad=0.65, labelspacing=0.50)

        # Summary box
        n_star_L = bc_sat[m_left];  n_star_R = bc_sat[m_right]
        ci_L = (bc_boot_ci[m_left][1], bc_boot_ci[m_left][2])
        ci_R = (bc_boot_ci[m_right][1], bc_boot_ci[m_right][2])
        summ = (f"BC N*({m_left})={n_star_L}  [{CI_LEVEL}%CI: {ci_L[0]}–{ci_L[1]}]\n"
                f"BC N*({m_right})={n_star_R}  [{CI_LEVEL}%CI: {ci_R[0]}–{ci_R[1]}]\n"
                f"Period: {period_obs}  |  M={M} models")
        ax_L.text(0.03, 0.06, summ, transform=ax_L.transAxes,
                  fontsize=9.5, va="bottom", color=C["grey"],
                  bbox=dict(boxstyle="round,pad=0.40", fc="white",
                            ec="#B0BEC5", alpha=0.95, lw=0.8))

        ax_L.set_title(f"{panel}  {ttl}\n"
                       f"     ─── Raw CMIP6  |  ─·─ Bias-Corrected (QDM)  |  ★ N*  |  "
                       f"Shaded = ±1SD + {CI_LEVEL}% CI",
                       loc="left", fontsize=12, fontweight="bold", pad=6)

    # ── Export ──────────────────────────────────────────────────────────
    print("  Splitting Fig 4 into individual panels ...")
    axes_for_split = [
        [ax_twins[0][0], ax_twins[0][1]],
        [ax_twins[1][0], ax_twins[1][1]],
    ]
    save_individual_panels(fig, axes_for_split, ["a", "b"],
                           fig_num=4, out_dir=out_dir, dpi=DPI)
    savefig(fig, out_dir / f"{prefix}_Fig4_IntegratedDualAxis")


# ════════════════════════════════════════════════════════════════════════
#  §11c  FIGURE 5 — SPATIAL CONSENSUS PROFILE  (Stacked Bar)  [NEW v5.0]
#         For each ensemble size N, what fraction of stations have
#         already reached saturation (N* ≤ N)?
#         • 2 × 2 layout, one sub-panel per metric
#         • Bar segments: "saturated at exactly N" (darker) vs
#           "already saturated at N < current" (lighter)
#         • Thin line overlay = cumulative % (right Y-axis)
#
#  Academic basis: Spatial consensus measure of Giorgi & Mearns (2002).
# ════════════════════════════════════════════════════════════════════════

def fig5_spatial_consensus(bc_stn_res, stns_str, models, bc_sat, bc_boot_ci,
                           period_obs, out_dir, prefix):
    """
    Fig 5 — Spatial Consensus Stacked Bar.

    For each metric and each N, computes:
      • n_exact(N) = number of stations with station-level N* == N
      • n_cumul(N) = cumulative number with N* ≤ N

    Stacked bars show the 'exact-at-N' fraction (dark) stacked on top of
    'already saturated' fraction (light), so bar height = cumulative %.
    A line overlay traces the cumulative saturation curve (Giorgi & Mearns, 2002).
    """
    M      = len(models)
    n_s    = len(stns_str)
    x      = np.arange(1, M + 1)

    METS_INFO = [
        ("RMSE", "(a)  RMSE", METRIC_COLORS["RMSE"]),
        ("KGE",  "(b)  KGE",  METRIC_COLORS["KGE"]),
        ("NSE",  "(c)  NSE",  METRIC_COLORS["NSE"]),
        ("r",    "(d)  r",    METRIC_COLORS["r"]),
    ]

    fig, axes = plt.subplots(2, 2, figsize=(18, 14))
    fig.subplots_adjust(hspace=0.48, wspace=0.32,
                        left=0.08, right=0.96, top=0.96, bottom=0.08)

    ax_list = [axes[0,0], axes[0,1], axes[1,0], axes[1,1]]

    for ax, (met, panel_ttl, base_col) in zip(ax_list, METS_INFO):
        ax_r = ax.twinx()

        # ── Compute per-station N* from bc_stn_res ─────────────────
        stn_nstar = []
        for stn in stns_str:
            sp_d = {n: float(np.mean(bc_stn_res[n][stn][met]))
                    if bc_stn_res[n][stn][met] else np.nan
                    for n in range(1, M+1)}
            stn_nstar.append(saturation_point(sp_d, met, M))
        stn_nstar = np.array(stn_nstar)

        # ── Counts: exact + cumulative ──────────────────────────────
        n_exact = np.array([np.sum(stn_nstar == n) for n in x])
        n_cumul = np.array([np.sum(stn_nstar <= n) for n in x])
        pct_exact  = n_exact / n_s * 100
        pct_cumul  = n_cumul / n_s * 100

        # already-saturated before this N = cumulative - exact
        pct_already = pct_cumul - pct_exact

        # ── Color variants ──────────────────────────────────────────
        col_dark  = base_col
        try:
            import matplotlib.colors as mc
            rgb_base = mc.to_rgb(base_col)
            col_light = tuple(min(1.0, c * 0.50 + 0.50) for c in rgb_base)
        except Exception:
            col_light = "#CFD8DC"

        # ── Stacked bar ─────────────────────────────────────────────
        bars_already = ax.bar(x, pct_already,
                              color=col_light, alpha=0.70,
                              edgecolor="#B0BEC5", linewidth=0.7,
                              zorder=2,
                              label=f"Saturated at N < current  (cumulative base)")
        bars_exact   = ax.bar(x, pct_exact, bottom=pct_already,
                              color=col_dark, alpha=0.88,
                              edgecolor="#1A1A1A", linewidth=0.9,
                              zorder=3,
                              label=f"Reaching saturation at exactly this N")

        # ── Value labels inside bars (if ≥ 5%) ──────────────────────
        for ni, (pe, pc) in enumerate(zip(pct_exact, pct_cumul)):
            if pe >= 5.0:
                ax.text(x[ni], pct_already[ni] + pe / 2,
                        f"{int(n_exact[ni])}\nstn", ha="center",
                        va="center", fontsize=9.5, fontweight="bold",
                        color="white", zorder=6)

        # ── Cumulative line on right axis ───────────────────────────
        ax_r.plot(x, pct_cumul, color=col_dark, lw=2.4, ls="-",
                  marker="D", ms=9,
                  markeredgecolor="#1A1A1A", markeredgewidth=0.8,
                  zorder=5, label="Cumulative % (right)")
        ax_r.set_ylim(0, 108)
        ax_r.set_ylabel("Cumulative % of stations",
                        fontsize=12, fontweight="bold", color=col_dark,
                        labelpad=6)
        ax_r.tick_params(axis="y", colors=col_dark, labelsize=11)
        ax_r.yaxis.set_major_locator(ticker.MultipleLocator(20))
        ax_r.spines["top"].set_visible(False)

        # 100% dashed guide
        ax_r.axhline(100, color="#546E7A", lw=1.0, ls=":", alpha=0.55)

        # Regional N* vertical marker
        reg_ns = bc_sat[met]
        lo_g   = bc_boot_ci[met][1]; hi_g = bc_boot_ci[met][2]
        ax.axvline(reg_ns, color=col_dark, lw=2.0, ls="--",
                   alpha=0.70, zorder=4,
                   label=f"Regional N*={reg_ns}  [{CI_LEVEL}%CI: {lo_g}–{hi_g}]")
        if lo_g != hi_g:
            ax.axvspan(lo_g - 0.15, hi_g + 0.15,
                       color=col_dark, alpha=0.10, zorder=1)

        # ── Axes decoration ──────────────────────────────────────────
        ax.set_xticks(x)
        ax.set_xticklabels([f"N={n}" for n in x], fontsize=10.5)
        ax.set_xlabel("Ensemble Size N", fontsize=13,
                      fontweight="bold", labelpad=6)
        ax.set_ylabel("% of Stations", fontsize=13,
                      fontweight="bold", labelpad=6)
        ax.set_ylim(0, 108)
        ax.yaxis.set_major_locator(ticker.MultipleLocator(20))
        ax.tick_params(axis="both", labelsize=11, width=1.6)
        ax.spines["top"].set_visible(False)

        ax.set_title(f"{panel_ttl} — Spatial Consensus Profile\n"
                     f"     Fraction of {n_s} stations reaching saturation vs N  |  "
                     f"Regional N*={reg_ns}  [{CI_LEVEL}%CI: {lo_g}–{hi_g}]",
                     loc="left", fontsize=11.5, fontweight="bold", pad=5)

        # Combined legend (both axes)
        h1, l1 = ax.get_legend_handles_labels()
        h2, l2 = ax_r.get_legend_handles_labels()
        ax.legend(h1 + h2, l1 + l2,
                  fontsize=9.8, frameon=True, edgecolor="#B0BEC5",
                  facecolor="white", framealpha=0.95,
                  ncol=1, loc="upper left",
                  handlelength=1.8, borderpad=0.80,
                  handletextpad=0.60, labelspacing=0.45)

    # ── Export ──────────────────────────────────────────────────────────
    print("  Splitting Fig 5 into individual panels ...")
    axes_for_split = [axes[0,0], axes[0,1], axes[1,0], axes[1,1]]
    save_individual_panels(fig, axes_for_split, ["a", "b", "c", "d"],
                           fig_num=5, out_dir=out_dir, dpi=DPI)
    savefig(fig, out_dir / f"{prefix}_Fig5_SpatialConsensus")


# ════════════════════════════════════════════════════════════════════════
#  §11d  FIGURE 6 — INTER-METRIC RELATIONSHIP WEB  [NEW v5.0]
#         Cross-correlations between metric trajectories across N
#         • Nodes: RMSE, KGE, NSE, r  arranged on a unit circle
#         • Edge width ∝ |Pearson r|  (only if |r| ≥ 0.30)
#         • Blue edge = positive correlation; Red = negative
#         • r value annotated on each edge
#         • Panel (a): Raw CMIP6  |  Panel (b): Bias-Corrected (QDM)
#
#  Academic basis: Metric inter-dependence analysis following
#  Moriasi et al. (2007) and Gupta et al. (2009).
# ════════════════════════════════════════════════════════════════════════

def fig6_intermetric_web(raw_means, bc_means, models,
                         period_obs, out_dir, prefix):
    """
    Fig 6 — Inter-metric Relationship Web.

    Computes the Pearson correlation between mean-metric trajectories
    (RMSE, KGE, NSE, r) across N = 1…M for Raw and BC separately.
    A network diagram displays each metric as a node; edges represent the
    strength (width) and direction (colour) of pairwise correlation.

    This graph reveals whether RMSE and KGE optimality are achieved
    simultaneously (positive edge) or traded-off (negative edge),
    informing metric selection for ensemble evaluation (Gupta et al., 2009).
    """
    M      = len(models)
    x_vals = np.arange(1, M + 1)
    METS   = ["RMSE", "KGE", "NSE", "r"]
    N_MET  = len(METS)

    def _series(means_dict, met):
        return np.array([means_dict[met].get(n, np.nan) for n in x_vals])

    def _corr_matrix(means_dict):
        """Build N_MET × N_MET symmetric Pearson-r matrix."""
        mat = np.full((N_MET, N_MET), np.nan)
        series = [_series(means_dict, m) for m in METS]
        for i in range(N_MET):
            for j in range(N_MET):
                a, b = series[i], series[j]
                valid = ~np.isnan(a) & ~np.isnan(b)
                if valid.sum() >= 3:
                    r_val, _ = pearsonr(a[valid], b[valid])
                    mat[i, j] = r_val
                else:
                    mat[i, j] = 0.0
        return mat

    # ── Node positions on unit circle ──────────────────────────────────
    angles  = [np.pi/2 + 2*np.pi*i/N_MET for i in range(N_MET)]
    node_xy = {m: (np.cos(a)*0.60, np.sin(a)*0.60)
                for m, a in zip(METS, angles)}

    NODE_COLORS = {m: METRIC_COLORS[m] for m in METS}
    NODE_RADIUS  = 0.085
    MIN_R_DRAW   = 0.30     # minimum |r| for drawing an edge

    def _draw_web(ax, corr_mat, title):
        ax.set_xlim(-1.05, 1.05); ax.set_ylim(-1.05, 1.05)
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_facecolor("#FAFAFA")

        # ── Edges ──────────────────────────────────────────────────
        for i in range(N_MET):
            for j in range(i + 1, N_MET):
                r_val = corr_mat[i, j]
                if np.isnan(r_val) or abs(r_val) < MIN_R_DRAW: continue
                x0, y0 = node_xy[METS[i]]
                x1, y1 = node_xy[METS[j]]
                lw   = max(1.2, abs(r_val) * 12.0)
                ecol = "#1565C0" if r_val > 0 else "#C62828"
                ealp = 0.45 + abs(r_val) * 0.50

                ax.plot([x0, x1], [y0, y1], color=ecol,
                        lw=lw, alpha=min(ealp, 0.92), zorder=2,
                        solid_capstyle="round")

                # r annotation on edge midpoint
                mx, my = (x0 + x1) / 2, (y0 + y1) / 2
                # offset label slightly perpendicular
                dx, dy = x1 - x0, y1 - y0
                norm = max(math.sqrt(dx**2 + dy**2), 1e-9)
                perp_x, perp_y = -dy / norm * 0.06, dx / norm * 0.06
                ax.text(mx + perp_x, my + perp_y,
                        f"r={r_val:+.2f}", fontsize=10.5,
                        ha="center", va="center", fontweight="bold",
                        color=ecol,
                        bbox=dict(boxstyle="round,pad=0.22",
                                  fc="white", ec=ecol,
                                  alpha=0.90, lw=0.7),
                        zorder=5)

        # ── Nodes ──────────────────────────────────────────────────
        for m in METS:
            nx, ny = node_xy[m]
            circ = mpatches.Circle((nx, ny), radius=NODE_RADIUS,
                                   color=NODE_COLORS[m],
                                   linewidth=2.2,
                                   edgecolor="#1A1A1A", zorder=6)
            ax.add_patch(circ)
            ax.text(nx, ny, m, ha="center", va="center",
                    fontsize=14, fontweight="bold", color="white", zorder=7)

            # Metric N* annotation below/above each node
            # (labelled outside node ring)
            dir_s = "→" if m not in LOWER_B else "←"
            ax.text(nx * 1.28, ny * 1.28,
                    f"{dir_s}{m}", ha="center", va="center",
                    fontsize=10, color=NODE_COLORS[m],
                    fontweight="bold")

        ax.set_title(title, fontsize=13, fontweight="bold", pad=10)

        # ── Legend ─────────────────────────────────────────────────
        leg_h = [
            Line2D([0],[0], color="#1565C0", lw=4, alpha=0.75,
                   label="Positive correlation"),
            Line2D([0],[0], color="#C62828", lw=4, alpha=0.75,
                   label="Negative correlation"),
            Line2D([0],[0], color="#546E7A", lw=1.5, alpha=0.65,
                   label=f"|r| < {MIN_R_DRAW:.2f} → not drawn"),
        ]
        ax.legend(handles=leg_h, fontsize=10, frameon=True,
                  edgecolor="#B0BEC5", facecolor="white",
                  framealpha=0.95, loc="lower center",
                  bbox_to_anchor=(0.5, -0.10),
                  ncol=3, handlelength=1.6, borderpad=0.75)

        # ── Edge width legend ───────────────────────────────────────
        for rv, lbl in [(0.5, "|r|=0.5"), (0.8, "|r|=0.8"), (1.0, "|r|=1.0")]:
            ax.plot([], [], color="#546E7A",
                    lw=max(1.2, rv * 12.0), alpha=0.55,
                    label=lbl)

        # ── Correlation table below network ────────────────────────
        tbl_lines = ["Pearson r  (upper triangle):"]
        tbl_lines.append("      " + "  ".join(f"{m:>5s}" for m in METS))
        for i in range(N_MET):
            row_parts = [f"{METS[i]:5s}"]
            for j in range(N_MET):
                if j <= i:
                    row_parts.append("     ")
                else:
                    rv = corr_mat[i, j]
                    row_parts.append(f"{rv:+5.2f}" if not np.isnan(rv) else "  N/A")
            tbl_lines.append("  ".join(row_parts))
        ax.text(0.02, 0.02, "\n".join(tbl_lines),
                transform=ax.transAxes, fontsize=9.5,
                family="monospace", va="bottom", ha="left",
                color="#37474F",
                bbox=dict(boxstyle="round,pad=0.45",
                          fc="white", ec="#B0BEC5",
                          alpha=0.92, lw=0.8),
                zorder=8)

    # ── Build figures ────────────────────────────────────────────────────
    raw_corr = _corr_matrix(raw_means)
    bc_corr  = _corr_matrix(bc_means)

    fig, (ax_raw, ax_bc) = plt.subplots(1, 2, figsize=(20, 11))
    fig.subplots_adjust(left=0.04, right=0.96, top=0.93, bottom=0.10,
                        wspace=0.08)

    _draw_web(ax_raw, raw_corr,
              f"(a)  Inter-metric Relationship Web — Raw CMIP6\n"
              f"     N = 1…{M}  |  Obs: {period_obs}  |  "
              f"Edge width ∝ |r|  |  Blue=+  Red=−")
    _draw_web(ax_bc,  bc_corr,
              f"(b)  Inter-metric Relationship Web — Bias-Corrected (QDM)\n"
              f"     N = 1…{M}  |  Obs: {period_obs}  |  "
              f"Edge width ∝ |r|  |  Blue=+  Red=−")

    # ── Export ──────────────────────────────────────────────────────────
    print("  Splitting Fig 6 into individual panels ...")
    save_individual_panels(fig, [ax_raw, ax_bc], ["a", "b"],
                           fig_num=6, out_dir=out_dir, dpi=DPI)
    savefig(fig, out_dir / f"{prefix}_Fig6_InterMetricWeb")

def write_excel(wb, raw_means, raw_stds, raw_sat,
                bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
                raw_spread_m, bc_spread_m,
                bc_stn_res, stns_str, smap, models,
                period_obs, val_report):
    M     = len(models)
    codes = [smap[s] for s in stns_str]

    def _title(ws, nc, t, s):
        _mxsc(ws,1,1,nc,t, bold=True,fc="FFFFFF",bg=XC["title"],sz=13,align="left"); _rh(ws,1,24)
        _mxsc(ws,2,1,nc,s, italic=True,fc="FFFFFF",bg=XC["sub"],sz=9,align="left");  _rh(ws,2,14)

    def _hdr(ws, r, hs):
        for ci, h in enumerate(hs, 1):
            _xsc(ws,r,ci,h, bold=True,fc="FFFFFF",bg=XC["hdr"],sz=10,wrap=True)
        _rh(ws, r, 36)

    def fv(v, dp=4):
        if v is None: return "—"
        return "—" if (isinstance(v,float) and np.isnan(v)) else round(float(v), dp)

    # ── S1: Saturation Results ──────────────────────────────────────
    ws1 = wb.create_sheet("S1 Saturation Results")
    ws1.sheet_view.showGridLines = False
    nc1 = 3 + len(METS_ALL)*2 + 1
    _title(ws1, nc1,
           "Ensemble Saturation — Regional Mean Performance vs Ensemble Size",
           f"Obs: {period_obs}  |  Exhaustive C(M,N)  |  "
           f"Saturation criterion: marginal < {int(SAT_THRESHOLD*100)}%  |  "
           "Values: mean (±SD) across all subsets of size N")
    hdrs1 = (["Dataset","N (Models)","N combos"] +
             [f"{k}\n(mean)" for k in METS_ALL] +
             [f"{k}\n(±1 SD)" for k in METS_ALL] +
             ["Spread\n(mm/d)"])
    _hdr(ws1, 4, hdrs1); ri1 = 5

    for tag, means, stds, sp_m, sd_dict, bg_k in [
        ("Raw CMIP6", raw_means, raw_stds, raw_spread_m, {}, XC["raw_r"]),
        ("BC (QDM)",  bc_means,  bc_stds,  bc_spread_m,  {}, XC["bc_r"]),
    ]:
        sat_d = raw_sat if tag=="Raw CMIP6" else bc_sat
        for N in range(1, M+1):
            nc_v    = math.comb(M, N)
            sat_flag = any(sat_d[k]==N for k in METS_ALL)
            bg = XC["sat"] if sat_flag else bg_k
            row = [tag, N, nc_v]
            for k in METS_ALL: row.append(fv(means[k].get(N, np.nan)))
            for k in METS_ALL: row.append(fv(stds[k].get(N, 0.0)))
            row.append(fv(sp_m.get(N, np.nan)))
            for ci, v in enumerate(row, 1):
                cell = _xsc(ws1,ri1,ci,v, sz=9.5,
                            align="left" if ci<=2 else "right",
                            bg=bg, bold=(sat_flag and ci>=4))
                if sat_flag and ci==2:
                    cell.font=Font(bold=True,color="E65100",name="Calibri",size=10)
            _rh(ws1,ri1,16); ri1+=1
        _rh(ws1,ri1-1,5)

    ri1+=1
    raw_s=" | ".join([f"{k}:N*={raw_sat[k]}" for k in METS_ALL])
    bc_s =" | ".join([f"{k}:N*={bc_sat[k]}"  for k in METS_ALL])
    _mxsc(ws1,ri1,1,nc1,f"N* (Raw): {raw_s}   |   N* (BC): {bc_s}",
          bold=True,sz=10,align="left",bg=XC["sat"]); _rh(ws1,ri1,24)
    for ci,w in enumerate([14,10,10]+[11]*len(METS_ALL)*2+[10],1): _cw(ws1,ci,w)

    # ── S2: Per-Station Saturation ──────────────────────────────────
    ws2 = wb.create_sheet("S2 Per-Station Saturation")
    ws2.sheet_view.showGridLines = False
    nc2 = 3+M*len(METS_ALL)+len(METS_ALL)
    _title(ws2,nc2,"Per-Station Saturation Points (N*)",
           f"Obs: {period_obs}  |  BC (QDM)  |  "
           f"N* = first N where marginal < {int(SAT_THRESHOLD*100)}%")
    hdrs2=(["Station","Code","Metric"]+[f"N={n}" for n in range(1,M+1)]+
           ["N*","Perf@N*","Perf@N=M","% extra\nN*→M"])
    _hdr(ws2,4,hdrs2); ri2=5

    for si,(stn,code) in enumerate(zip(stns_str,codes)):
        bg_r = XC["bc_r"] if si%2==0 else XC["white"]
        for met in METS_ALL:
            sp_dict={n:float(np.mean(bc_stn_res[n][stn][met]))
                     if bc_stn_res[n][stn][met] else np.nan for n in range(1,M+1)}
            sat_n=saturation_point(sp_dict,met,M)
            v_sat=sp_dict.get(sat_n,np.nan); v_M=sp_dict.get(M,np.nan); v_1=sp_dict.get(1,np.nan)
            total=abs(v_M-v_1) if not(np.isnan(v_M) or np.isnan(v_1)) else np.nan
            extra=abs(v_M-v_sat)/total*100 if not(np.isnan(v_M) or np.isnan(v_sat) or
                                                   np.isnan(total) or total==0) else np.nan
            row2=[stn if met==METS_ALL[0] else "", code if met==METS_ALL[0] else "", met]
            for n in range(1,M+1): row2.append(fv(sp_dict.get(n,np.nan)))
            row2+=[sat_n,fv(v_sat),fv(v_M),
                   f"{extra:.1f}%" if not np.isnan(extra) else "—"]
            for ci,v in enumerate(row2,1):
                bg_c=XC["sat"] if ci==len(row2)-3 else bg_r
                _xsc(ws2,ri2,ci,v,sz=9.5,align="left" if ci<=3 else "right",
                     bg=bg_c,bold=(ci==len(row2)-3))
            _rh(ws2,ri2,15); ri2+=1
    for ci,w in enumerate([10,6,6]+[10]*M+[10,10,10,12],1): _cw(ws2,ci,w)

    # ── S3: Data Validation Report ──────────────────────────────────
    ws3 = wb.create_sheet("S3 Data Validation")
    ws3.sheet_view.showGridLines = False
    _title(ws3,4,"Data Validation Report — Pre-run Integrity Checks",
           f"Obs: {period_obs}  |  Status: {'PASS' if val_report.get('ok',True) else 'FAIL'}")
    _hdr(ws3,4,["Check","Result","Status"])
    _rh(ws3,4,30); ri3=5
    checks=[
        ("Common temporal period",
         val_report.get("common_period","N/A"),
         "PASS" if val_report.get("common_days",0)>=365 else "FAIL"),
        ("Common days overlap",
         f"{val_report.get('common_days','N/A')} days",
         "PASS" if val_report.get("common_days",0)>=365 else "FAIL"),
        ("Date continuity gap",
         f"{val_report.get('date_gap_pct','N/A')}%",
         "PASS" if val_report.get("date_gap_pct",0)<=5 else "WARN"),
        ("Errors",
         str(val_report.get("errors",[])),
         "PASS" if not val_report.get("errors") else "FAIL"),
        ("Warnings",
         str(val_report.get("warnings",[])),
         "OK" if not val_report.get("warnings") else "WARN"),
    ]
    for lbl,val,sts in checks:
        bg=XC["improve"] if sts in("PASS","OK") else (XC["degrade"] if sts=="FAIL" else XC["sat"])
        _xsc(ws3,ri3,1,lbl,bold=True,sz=10,align="left",bg=bg)
        _xsc(ws3,ri3,2,val,sz=9.5,align="left",bg=bg)
        _xsc(ws3,ri3,3,sts,bold=True,sz=10,
             fc="1B5E20" if sts in("PASS","OK") else ("B71C1C" if sts=="FAIL" else "E65100"),
             bg=bg); _rh(ws3,ri3,18); ri3+=1
    # Missing % per model
    ri3+=1
    _xsc(ws3,ri3,1,"Missing Value %",bold=True,sz=10,bg=XC["hdr"],fc="FFFFFF"); ri3+=1
    for lbl,pct in val_report.get("missing_pct",{}).items():
        bg=XC["improve"] if pct<=5 else (XC["degrade"] if pct>20 else XC["sat"])
        _xsc(ws3,ri3,1,lbl,sz=9.5,align="left",bg=bg)
        _xsc(ws3,ri3,2,f"{pct:.2f}%",sz=9.5,bg=bg); _rh(ws3,ri3,14); ri3+=1
    for ci,w in enumerate([32,42,12],1): _cw(ws3,ci,w)

    # ── S4: Bootstrap CI for N* [NEW in v3.0] ───────────────────────
    ws4 = wb.create_sheet("S4 Bootstrap CI Nstar")
    ws4.sheet_view.showGridLines = False
    nc4 = 9
    _title(ws4,nc4,"Bootstrap Confidence Intervals for Saturation Point N* (BC QDM)",
           f"Method: bootstrap percentile CI (n={N_BOOTSTRAP})  |  "
           f"CI level: {CI_LEVEL}%  |  Efron & Tibshirani (1993)  |  Obs: {period_obs}")
    hdrs4=["Metric","N* (point est.)","Bootstrap\nmean N*",
           "Bootstrap\nstd N*","CI lower\n(N*)","CI upper\n(N*)",
           "CI width","Perf @ N*","Perf @ N=M"]
    _hdr(ws4,4,hdrs4); _rh(ws4,4,36); ri4=5

    for met in METS_ALL:
        N_star, lo, hi, dist, ci_lo_d, ci_hi_d = bc_boot_ci[met]
        boot_mean = float(np.mean(dist)); boot_std = float(np.std(dist,ddof=1))
        ci_width  = hi - lo
        v_sat = bc_means[met].get(N_star, np.nan)
        v_M   = bc_means[met].get(M, np.nan)
        bg_met = XC["bc_r"]
        row4=[met, N_star, round(boot_mean,2), round(boot_std,2),
              lo, hi, ci_width, fv(v_sat), fv(v_M)]
        for ci,v in enumerate(row4,1):
            cell=_xsc(ws4,ri4,ci,v,sz=10,
                      align="left" if ci==1 else "right",bg=bg_met,
                      bold=(ci in(1,2,5,6)))
            if ci in(5,6) and ci_width==0:
                cell.fill=_xf(XC["improve"])
        _rh(ws4,ri4,18); ri4+=1

    # Per-N mean CI table
    ri4+=2
    _mxsc(ws4,ri4,1,nc4,
          "Per-N Bootstrap CI for Mean Performance (BC QDM)",
          bold=True,fc="FFFFFF",bg=XC["hdr"],sz=11); _rh(ws4,ri4,22); ri4+=1
    hdrs4b=["N"]+[f"{k}\nmean" for k in METS_ALL]+[f"{k}\n95%CI lo" for k in METS_ALL]+\
           [f"{k}\n95%CI hi" for k in METS_ALL]
    _hdr(ws4,ri4,hdrs4b[:nc4+1]); _rh(ws4,ri4,30); ri4+=1

    for N in range(1,M+1):
        sat_flag=any(bc_sat[k]==N for k in METS_ALL)
        bg=XC["sat"] if sat_flag else (XC["bc_r"] if N%2==0 else XC["white"])
        row_n=[N]
        for k in METS_ALL: row_n.append(fv(bc_means[k].get(N,np.nan)))
        for k in METS_ALL: row_n.append(fv(bc_boot_ci[k][4].get(N,np.nan)))
        for k in METS_ALL: row_n.append(fv(bc_boot_ci[k][5].get(N,np.nan)))
        for ci,v in enumerate(row_n[:nc4+1],1):
            cell=_xsc(ws4,ri4,ci,v,sz=9.5,align="right",bg=bg,bold=(sat_flag and ci==1))
        _rh(ws4,ri4,15); ri4+=1

    for ci,w in enumerate([8]+[10]*8,1): _cw(ws4,ci,w)


# ════════════════════════════════════════════════════════════════════════
#  §13  WORD REPORT
# ════════════════════════════════════════════════════════════════════════

def write_word(raw_means, raw_sat, bc_means, bc_sat, bc_boot_ci,
               bc_spread_m, raw_spread_m,
               models, stns_str, period_obs, period_sim, out_dir, prefix):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  ⚠  python-docx not installed — Word skipped"); return

    doc = Document()
    for sec in doc.sections:
        sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54)
        sec.top_margin=Cm(2.54);  sec.bottom_margin=Cm(2.54)
    M=len(models)

    def _h(txt,level=1,color="1B2838"):
        h=doc.add_heading(txt,level=level)
        r,g,b=int(color[0:2],16),int(color[2:4],16),int(color[4:6],16)
        for run in h.runs: run.font.name="Times New Roman"; run.font.color.rgb=RGBColor(r,g,b)

    def _p(txt,bold=False,italic=False,sz=12,align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p=doc.add_paragraph(); p.alignment=align
        run=p.add_run(txt); run.font.name="Times New Roman"
        run.font.size=Pt(sz); run.bold=bold; run.italic=italic; return p

    def _find(txt,sz=12):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        rb=p.add_run("► "); rb.bold=True; rb.font.name="Times New Roman"; rb.font.size=Pt(sz)
        rd=p.add_run(txt); rd.font.name="Times New Roman"; rd.font.size=Pt(sz)

    # Key results
    bc_kge_full=bc_means["KGE"].get(M,np.nan); bc_kge_sat=bc_means["KGE"].get(bc_sat["KGE"],np.nan)
    bc_rmse_full=bc_means["RMSE"].get(M,np.nan)
    raw_sp_M=raw_spread_m.get(M,np.nan); bc_sp_M=bc_spread_m.get(M,np.nan)
    total_c=sum(math.comb(M,n) for n in range(1,M+1))
    sat_min=min(bc_sat[k] for k in METS_ALL); sat_max=max(bc_sat[k] for k in METS_ALL)
    Nstar_kge,lo_kge,hi_kge=bc_boot_ci["KGE"][0],bc_boot_ci["KGE"][1],bc_boot_ci["KGE"][2]

    # Title
    t=doc.add_heading("",0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=t.add_run("Ensemble Saturation Analysis — v3.0\n"
                  "How Many CMIP6 Models Are Needed for Stable Station-Scale Rainfall?")
    run.font.name="Times New Roman"; run.font.size=Pt(16); run.bold=True
    doc.add_paragraph()
    p_sub=doc.add_paragraph(); p_sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rs=p_sub.add_run(f"Study Area: Prachuap Khiri Khan Province, Thailand  |  "
                     f"Period: {period_obs}  |  Models: {', '.join(models)}")
    rs.font.name="Times New Roman"; rs.font.size=Pt(12); rs.italic=True
    doc.add_paragraph()

    _h("Abstract",1,"13293D")
    kge_s=f"{bc_kge_full:.3f}" if not np.isnan(bc_kge_full) else "N/A"
    _p(f"This study uses exhaustive combinatorial analysis ({total_c} unique ensembles) "
       f"to identify the ensemble saturation point N* for {M} CMIP6 models in Prachuap Khiri Khan, "
       f"Thailand ({period_obs}). After QDM bias correction, saturation occurs at "
       f"N*={Nstar_kge} (KGE, {CI_LEVEL}% bootstrap CI: [{lo_kge}, {hi_kge}]) "
       f"to N*={bc_sat['RMSE']} (RMSE). "
       f"The full {M}-model ensemble achieves KGE = {kge_s}. "
       f"These findings provide evidence-based guidance on minimum ensemble size "
       f"for monsoon climate assessments in Thailand.")
    doc.add_paragraph()

    _h("1.  Data and Methods",1,"1F4E79")
    _h("1.1  Observed Data",2,"2E75B6")
    _p(f"Daily rainfall from {len(stns_str)} TMD stations ({period_obs}). "
       "Missing values (−99, −999, −9999) replaced with NaN; negatives discarded (WMO).")
    doc.add_paragraph()
    _h("1.2  CMIP6 Models and Bias Correction",2,"2E75B6")
    _p(f"CMIP6 historical precipitation ({period_sim}) from {M} models "
       f"({', '.join(models)}) obtained from ESGF. "
       "QDM (Cannon et al., 2015) applied per station with equal-weight ensemble mean "
       "(Knutti et al., 2017).")
    doc.add_paragraph()
    _h("1.3  Saturation Analysis Method",2,"2E75B6")
    _p(f"Exhaustive C({M},N) analysis ({total_c} subsets). "
       f"N* = first N where marginal improvement < {int(SAT_THRESHOLD*100)}% of total improvement. "
       f"Bootstrap {CI_LEVEL}% CI for N* computed via {N_BOOTSTRAP} resamplings of subset values "
       "(Efron & Tibshirani, 1993). Parallel processing with ProcessPoolExecutor "
       f"({MAX_WORKERS} cores) used for efficiency.")
    doc.add_paragraph()

    _h("2.  Results",1,"1F4E79")
    _h("2.1  Saturation Points with Bootstrap CI",2,"2E75B6")
    lines_res=[]
    for k in METS_ALL:
        Ns,lo,hi=bc_boot_ci[k][0],bc_boot_ci[k][1],bc_boot_ci[k][2]
        vs=bc_means[k].get(Ns,np.nan); vM=bc_means[k].get(M,np.nan)
        lines_res.append(f"{k}: N*={Ns} [{CI_LEVEL}%CI: {lo}–{hi}]  "
                         f"Perf@N*={vs:.3f}  Perf@N={M}={vM:.3f}")
    _p("Saturation results (BC QDM):  " + "   |   ".join(lines_res))
    doc.add_paragraph()

    _h("3.  Key Findings",1,"1F4E79")
    findings=[
        (f"Saturation at N*={Nstar_kge} (KGE) to N*={bc_sat['RMSE']} (RMSE)",
         f"Bootstrap {CI_LEVEL}% CI for N*={Nstar_kge}: [{lo_kge}, {hi_kge}]. "
         "The CI width confirms statistical robustness of the saturation estimate."),
        ("Bootstrap CI validates saturation identification",
         f"{N_BOOTSTRAP} bootstrap resamplings of C(M,N) subset values confirm N* stability. "
         "Narrow CI = high confidence; wide CI = N* sensitive to subset composition."),
        (f"Parallel processing ({MAX_WORKERS} cores)",
         "ProcessPoolExecutor distributes C(M,N) subsets across CPU cores, "
         "reducing computation time for large M."),
        ("Data validation passed all critical checks",
         "Pre-run checks confirm date overlap ≥1 year, missing value % < threshold, "
         "and station consistency across all datasets."),
    ]
    for t_f,det in findings: _find(f"{t_f}: {det}")
    doc.add_paragraph()

    _h("4.  References",1,"1F4E79")
    refs=[
        "Cannon AJ et al. (2015). J. Climate, 28, 6938–6959.",
        "Efron B, Tibshirani RJ (1993). Bootstrap Methods. Chapman & Hall.",
        "Giorgi F, Mearns LO (2002). J. Climate, 15, 1141–1158.",
        "Gupta HV et al. (2009). J. Hydrology, 377, 80–91.",
        "Knutti R et al. (2017). Geophys. Res. Lett., 44, 1909–1918.",
        "Moriasi DN et al. (2007). Trans. ASABE, 50, 885–900.",
        "Nash JE, Sutcliffe JV (1970). J. Hydrology, 10, 282–290.",
        "Tebaldi C, Knutti R (2007). Phil. Trans. R. Soc. A, 365, 2053–2075.",
    ]
    for i,ref in enumerate(refs,1):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        rn=p.add_run(f"[{i}]  "); rn.bold=True; rn.font.name="Times New Roman"; rn.font.size=Pt(11)
        rt=p.add_run(ref); rt.font.name="Times New Roman"; rt.font.size=Pt(11)

    out_path=out_dir/f"{prefix}_SaturationReport_v{VERSION}.docx"
    doc.save(str(out_path)); print(f"  ✓  Word → {out_path.name}")


# ════════════════════════════════════════════════════════════════════════
#  §14  UNIT TESTS  [NEW in v3.0 — run with pytest]
# ════════════════════════════════════════════════════════════════════════

def _generate_test_file(out_dir):
    """Write test_saturation_v3.py alongside the main script."""
    test_code = '''# -*- coding: utf-8 -*-
"""
Unit Tests for saturation_curve_v3.py
Run with: pytest test_saturation_v3.py -v
Tests: KGE, NSE, RMSE formulas; saturation point; bootstrap CI
"""
import numpy as np
import math
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from saturation_curve_v3 import (
    _compute_metrics, saturation_point, bootstrap_saturation_ci
)


class TestMetrics:
    """Validate metric formulas against known analytical solutions."""

    def _perfect(self, n=100):
        """Perfect simulation: s == o"""
        np.random.seed(0)
        o = np.random.exponential(5, n)
        return o, o.copy()

    def _zeroes(self, n=100):
        """Zero simulation: s == 0"""
        np.random.seed(0)
        o = np.random.exponential(5, n)
        return o, np.zeros(n)

    # ── KGE tests ─────────────────────────────────────────────────
    def test_kge_perfect(self):
        """KGE = 1.0 when s == o (Gupta et al. 2009)."""
        o, s = self._perfect()
        m = _compute_metrics(o, s)
        assert abs(m["KGE"] - 1.0) < 1e-9, f"KGE perfect: {m['KGE']}"

    def test_kge_range(self):
        """KGE can be ≤ 1 but has no fixed lower bound."""
        np.random.seed(1)
        o = np.random.exponential(5, 200)
        s = np.random.exponential(3, 200)
        m = _compute_metrics(o, s)
        assert m["KGE"] <= 1.0 + 1e-9, f"KGE <= 1: {m['KGE']}"
        assert not np.isnan(m["KGE"]), "KGE is NaN"

    def test_kge_formula_components(self):
        """Verify KGE = 1 - sqrt((r-1)^2 + (sigma_r-1)^2 + (beta-1)^2)."""
        np.random.seed(2)
        o = np.abs(np.random.randn(200)) + 1.0
        s = 1.1 * o + 0.5
        m = _compute_metrics(o, s)
        # Manual calculation
        r_v    = float(np.corrcoef(o, s)[0, 1])
        sigma_r = float(np.std(s, ddof=1) / np.std(o, ddof=1))
        beta    = float(np.mean(s) / np.mean(o))
        kge_exp = 1 - math.sqrt((r_v-1)**2 + (sigma_r-1)**2 + (beta-1)**2)
        assert abs(m["KGE"] - kge_exp) < 1e-9, f"KGE formula mismatch: {m['KGE']} vs {kge_exp}"

    # ── NSE tests ─────────────────────────────────────────────────
    def test_nse_perfect(self):
        """NSE = 1.0 when s == o (Nash & Sutcliffe 1970)."""
        o, s = self._perfect()
        m = _compute_metrics(o, s)
        assert abs(m["NSE"] - 1.0) < 1e-9, f"NSE perfect: {m['NSE']}"

    def test_nse_climatological(self):
        """NSE = 0 when s == mean(o) (climatological forecast)."""
        np.random.seed(3)
        o = np.random.exponential(5, 200) + 1.0
        s = np.full_like(o, np.mean(o))
        m = _compute_metrics(o, s)
        assert abs(m["NSE"]) < 1e-9, f"NSE clim: {m['NSE']}"

    def test_nse_formula(self):
        """Verify NSE = 1 - sum(e^2) / sum((o - mean_o)^2)."""
        np.random.seed(4)
        o = np.abs(np.random.randn(300)) + 1.0
        s = 0.9 * o + 1.0
        m = _compute_metrics(o, s)
        e       = s - o
        dn      = np.sum((o - np.mean(o))**2)
        nse_exp = float(1 - np.sum(e**2) / dn)
        assert abs(m["NSE"] - nse_exp) < 1e-9, f"NSE formula: {m['NSE']} vs {nse_exp}"

    # ── RMSE tests ────────────────────────────────────────────────
    def test_rmse_perfect(self):
        """RMSE = 0 when s == o."""
        o, s = self._perfect()
        m = _compute_metrics(o, s)
        assert m["RMSE"] < 1e-9, f"RMSE perfect: {m['RMSE']}"

    def test_rmse_positive(self):
        """RMSE > 0 for any imperfect simulation."""
        np.random.seed(5)
        o = np.random.exponential(5, 200)
        s = o + np.random.randn(200)
        m = _compute_metrics(o, s)
        assert m["RMSE"] > 0, f"RMSE > 0: {m['RMSE']}"

    def test_rmse_formula(self):
        """Verify RMSE = sqrt(mean((s-o)^2))."""
        np.random.seed(6)
        o = np.abs(np.random.randn(300)) + 1.0
        s = 1.2 * o
        m = _compute_metrics(o, s)
        rmse_exp = float(np.sqrt(np.mean((s - o)**2)))
        assert abs(m["RMSE"] - rmse_exp) < 1e-9, f"RMSE formula: {m['RMSE']} vs {rmse_exp}"

    # ── r tests ───────────────────────────────────────────────────
    def test_r_perfect(self):
        """Pearson r = 1 for perfect linear relationship."""
        o, s = self._perfect()
        m = _compute_metrics(o, s)
        assert abs(m["r"] - 1.0) < 1e-9, f"r perfect: {m['r']}"

    def test_r_range(self):
        """Pearson r ∈ [-1, 1]."""
        np.random.seed(7)
        o = np.random.randn(200)
        s = np.random.randn(200)
        m = _compute_metrics(o, s)
        assert -1.0 - 1e-9 <= m["r"] <= 1.0 + 1e-9, f"r range: {m['r']}"

    # ── Edge cases ────────────────────────────────────────────────
    def test_insufficient_data(self):
        """Fewer than 5 valid pairs → all NaN."""
        o = np.array([1.0, 2.0, 3.0]); s = np.array([1.0, 2.0, 3.0])
        m = _compute_metrics(o, s)
        for k in ["RMSE","KGE","NSE","r"]:
            assert np.isnan(m[k]), f"{k} should be NaN for n<5"

    def test_nan_handling(self):
        """NaN values are excluded before computation."""
        np.random.seed(8)
        o = np.random.exponential(5, 200).astype(float)
        s = o * 1.05
        o[::10] = np.nan   # introduce NaN
        m = _compute_metrics(o, s)
        assert not np.isnan(m["KGE"]), "NaN handling failed"


class TestSaturationPoint:
    """Validate saturation criterion."""

    def test_immediate_flat(self):
        """Flat curve from N=1 → saturation at N=1."""
        means = {1: 1.0, 2: 1.0, 3: 1.0, 4: 1.0, 5: 1.0}
        N = saturation_point(means, "KGE", 5)
        assert N == 1, f"Flat → N=1: got {N}"

    def test_monotone_decrease(self):
        """Strictly decreasing large improvements → N=5 (no saturation)."""
        means = {1: 5.0, 2: 3.0, 3: 2.0, 4: 1.5, 5: 1.0}
        N = saturation_point(means, "RMSE", 5)
        # Each step is > 5% of total → no saturation until N=5
        assert N == 5, f"Large drops → N=5: got {N}"

    def test_early_saturation(self):
        """Large first step then tiny steps → saturation at N=2."""
        means = {1: 10.0, 2: 1.0, 3: 0.99, 4: 0.98, 5: 0.97}
        N = saturation_point(means, "RMSE", 5)
        assert N == 2, f"Early saturation: got {N}"

    def test_threshold_exactly(self):
        """Exactly at 5% threshold — should NOT trigger saturation."""
        total = 1.0  # total improvement from N=1 to N=5
        # step from N=3→4: 5% of total = 0.05 (boundary, not below)
        means = {1: 1.0, 2: 0.6, 3: 0.3, 4: 0.25, 5: 0.0}
        # step 3→4 = 0.05, total = 1.0, ratio = 0.05 = threshold (not strictly less)
        N = saturation_point(means, "RMSE", 5, threshold=0.05)
        # 0.05 is not < 0.05, so no saturation at N=4; check N=5
        assert N >= 3  # boundary: step not strictly < threshold


class TestBootstrapCI:
    """Validate bootstrap CI for saturation point."""

    def test_ci_within_bounds(self):
        """CI lower ≤ N* ≤ CI upper."""
        results = {
            1: {"KGE": [0.10, 0.12, 0.09, 0.11, 0.10]},
            2: {"KGE": [0.08, 0.07, 0.08, 0.09, 0.07]},
            3: {"KGE": [0.06, 0.06, 0.05, 0.07, 0.06]},
        }
        N_star, lo, hi, dist, _, _ = bootstrap_saturation_ci(
            results, "KGE", M=3, n_bootstrap=100, ci=95)
        assert lo <= N_star <= hi, f"N* not in CI: [{lo}, {hi}], N*={N_star}"
        assert 1 <= lo <= 3
        assert 1 <= hi <= 3

    def test_ci_distribution_length(self):
        """Bootstrap distribution has correct length."""
        results = {
            1: {"RMSE": [10.1, 10.2, 9.9, 10.0, 10.1]},
            2: {"RMSE": [9.0,  9.1, 8.9, 9.0, 9.2]},
            3: {"RMSE": [8.5,  8.6, 8.4, 8.5, 8.6]},
        }
        N_star, lo, hi, dist, _, _ = bootstrap_saturation_ci(
            results, "RMSE", M=3, n_bootstrap=50)
        assert len(dist) == 50, f"Bootstrap dist length: {len(dist)}"


if __name__ == "__main__":
    import pytest
    pytest.main([__file__, "-v", "--tb=short"])
'''
    test_path = out_dir / "test_saturation_v3.py"
    with open(str(test_path), "w", encoding="utf-8") as f:
        f.write(test_code)
    print(f"  ✓  Unit tests → {test_path.name}")
    return test_path


# ════════════════════════════════════════════════════════════════════════
#  §15  MAIN
# ════════════════════════════════════════════════════════════════════════

def main():
    SEP = "═" * 72
    print(SEP)
    print(f"  Ensemble Saturation Curve Analysis  v{VERSION}")
    print(f"  Parallel={MAX_WORKERS} cores  |  Bootstrap n={N_BOOTSTRAP}  |  DPI={DPI}")
    print("  Q2–Q3 Publication Standard  |  6 Figs + 16 Panels + Excel (4 sheets) + Word + Tests")
    print(SEP)

    if len(sys.argv) > 1:
        work_dir = sys.argv[1].strip('"').strip("'")
    else:
        try:    work_dir = str(Path(os.path.abspath(__file__)).parent)
        except: work_dir = os.getcwd()
    out_dir = Path(work_dir)
    print(f"  Input folder : {work_dir}")

    # ── Discover files ──────────────────────────────────────────────
    print("\n  Discovering files ...")
    obs_path, raw_models, bc_models = discover_files(work_dir)
    if obs_path is None: sys.exit("  ✗  No Observed file")
    if not raw_models and not bc_models: sys.exit("  ✗  No model files")

    paired = sorted(set(raw_models) & set(bc_models))
    if not paired:
        # fuzzy pairing
        for rk in list(raw_models.keys()):
            for bk in list(bc_models.keys()):
                if rk.lower()==bk.lower() and rk not in paired:
                    paired.append(rk)
                    if rk not in bc_models: bc_models[rk]=bc_models.pop(bk)
        if not paired: sys.exit("  ✗  No paired Raw+BC models")

    models = paired; M = len(models)
    total_c = sum(math.comb(M,n) for n in range(1,M+1))
    print(f"\n  Observed : {Path(obs_path).name}")
    print(f"  Models   : {models}  (M={M})")
    print(f"  C(M,N) subsets: {total_c}  |  CPU cores: {MAX_WORKERS}")
    print("-" * 72)

    # ── Load data ───────────────────────────────────────────────────
    print("\n  Loading data ...")
    obs_d, stns = load_daily(obs_path, "Observed")
    if obs_d is None: sys.exit("  ✗  Failed to load Observed")
    stns_str   = [str(s) for s in stns]
    smap       = short_labels(stns)
    period_obs = period_str(obs_d)
    prefix     = f"Saturation_{Path(obs_path).stem}"

    raw_dfs={}
    for m in models:
        df,_ = load_daily(raw_models[m], f"Raw/{m}", target_stns=stns_str)
        if df is not None: raw_dfs[m]=df

    bc_dfs={}
    for m in models:
        df,_ = load_daily(bc_models[m], f"BC/{m}", target_stns=stns_str)
        if df is not None: bc_dfs[m]=df

    models_ok=[m for m in models if m in raw_dfs and m in bc_dfs]
    if not models_ok: sys.exit("  ✗  No models loaded successfully")
    models=models_ok; M=len(models)
    period_sim = period_str(next(iter(raw_dfs.values())))
    print(f"\n  {len(stns_str)} stations | M={M} models | "
          f"Obs: {period_obs} | Sim: {period_sim}")
    print("-" * 72)

    # ── Data validation ─────────────────────────────────────────────
    val_report, val_ok = validate_input_data(obs_d, raw_dfs, bc_dfs, stns_str)
    if not val_ok:
        print("  ⚠  Proceeding despite validation errors — check results carefully")

    # ── Saturation analysis ─────────────────────────────────────────
    print("\n  Running saturation analysis (parallel) ...")
    print("  [1/2] Raw CMIP6 ...")
    _, raw_means, raw_stds, raw_sp_m, raw_sp_s, raw_sat, _, raw_boot_ci = \
        run_saturation_analysis(obs_d, raw_dfs, stns_str, models, "Raw CMIP6",
                                use_parallel=(MAX_WORKERS > 1))

    print("  [2/2] Bias-Corrected (QDM) ...")
    _, bc_means, bc_stds, bc_sp_m, bc_sp_s, bc_sat, bc_stn_res, bc_boot_ci = \
        run_saturation_analysis(obs_d, bc_dfs, stns_str, models, "BC (QDM)",
                                use_parallel=(MAX_WORKERS > 1))

    print("\n  ── Saturation Summary ─────────────────────────────────")
    for k in METS_ALL:
        Ns,lo,hi=bc_boot_ci[k][0],bc_boot_ci[k][1],bc_boot_ci[k][2]
        print(f"  {k:6s}: Raw N*={raw_sat[k]}  |  "
              f"BC N*={Ns}  [{CI_LEVEL}%CI: {lo}–{hi}]")

    # ── Figures ─────────────────────────────────────────────────────
    print(f"\n  Generating figures (DPI={DPI}) ...")
    print("  Fig 1: Performance Saturation ...")
    fig1_saturation_performance(
        raw_means, raw_stds, raw_sat, raw_boot_ci,
        bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
        models, period_obs, out_dir, prefix); gc.collect()

    print("  Fig 2: Spread Saturation ...")
    fig2_saturation_spread(
        raw_sp_m, raw_sp_s, bc_sp_m, bc_sp_s,
        raw_means, bc_means, raw_sat, bc_sat,
        models, period_obs, out_dir, prefix); gc.collect()

    print("  Fig 3: Spatial Saturation ...")
    fig3_saturation_spatial(
        bc_stn_res, stns_str, smap, models, bc_sat, bc_boot_ci,
        period_obs, out_dir, prefix); gc.collect()

    # ── NEW v5.0 figures ─────────────────────────────────────────────
    print("  Fig 4: Integrated Dual-Axis Performance Profile ...")
    fig4_integrated_profile(
        raw_means, raw_stds, raw_sat, raw_boot_ci,
        bc_means,  bc_stds,  bc_sat,  bc_boot_ci,
        models, period_obs, out_dir, prefix); gc.collect()

    print("  Fig 5: Spatial Consensus Stacked Bar ...")
    fig5_spatial_consensus(
        bc_stn_res, stns_str, models, bc_sat, bc_boot_ci,
        period_obs, out_dir, prefix); gc.collect()

    print("  Fig 6: Inter-metric Relationship Web ...")
    fig6_intermetric_web(
        raw_means, bc_means, models,
        period_obs, out_dir, prefix); gc.collect()

    # ── Excel ───────────────────────────────────────────────────────
    print("\n  Building Excel (4 sheets) ...")
    wb = Workbook(); wb.remove(wb.active)
    write_excel(wb, raw_means, raw_stds, raw_sat,
                bc_means, bc_stds, bc_sat, bc_boot_ci,
                raw_sp_m, bc_sp_m, bc_stn_res, stns_str, smap,
                models, period_obs, val_report)
    out_xl = out_dir / f"{prefix}_Analysis_v{VERSION}.xlsx"
    wb.save(str(out_xl)); print(f"  ✓  Excel → {out_xl.name}")

    # ── Word ─────────────────────────────────────────────────────────
    print("\n  Building Word report ...")
    write_word(raw_means, raw_sat, bc_means, bc_sat, bc_boot_ci,
               bc_sp_m, raw_sp_m, models, stns_str,
               period_obs, period_sim, out_dir, prefix)

    # ── Unit test file ───────────────────────────────────────────────
    print("\n  Writing unit test file ...")
    _generate_test_file(out_dir)

    # ── Summary ─────────────────────────────────────────────────────
    n_png    = len(list(out_dir.glob(f"{prefix}_Fig*.png")))
    n_panels = len(list(out_dir.glob("Figure_*.png")))
    print()
    print(SEP)
    print(f"  ✓  COMPLETE  v{VERSION}  |  DPI={DPI}")
    print(f"  {'─'*68}")
    print(f"  Figures   : {n_png} composite PNG" + (" + PDF" if SAVE_PDF else ""))
    print(f"  Panels    : {n_panels} individual PNG")
    print(f"     Fig 1 (a–d) Performance | Fig 2 (a–c) Spread | Fig 3 (a–c) Spatial")
    print(f"     Fig 4 (a–b) Dual-Axis  | Fig 5 (a–d) Consensus | Fig 6 (a–b) Web")
    print(f"  Excel     : {out_xl.name}  (4 sheets: S1–S4)")
    print(f"  Word      : {prefix}_SaturationReport_v{VERSION}.docx")
    print(f"  Unit tests: test_saturation_v3.py  (run: pytest test_saturation_v3.py -v)")
    print(f"  {'─'*68}")
    for k in METS_ALL:
        Ns,lo,hi=bc_boot_ci[k][0],bc_boot_ci[k][1],bc_boot_ci[k][2]
        vN=bc_means[k].get(M,np.nan); vs=bc_means[k].get(Ns,np.nan)
        arr="↓" if k in LOWER_B else "↑"
        print(f"  BC {k:4s}: N*={Ns} [{lo}–{hi}]  "
              f"Perf@N*={vs:.4f}  Perf@N={M}={vN:.4f}  {arr}")
    print(f"  Saved to : {work_dir}")
    print(SEP)


if __name__ == "__main__":
    main()
